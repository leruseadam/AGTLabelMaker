#!/usr/bin/env python3
"""
AGT Price Tag Transformer
--------------------------
This script reads and processes an Excel file, applies various filters,
generates dynamic DOCX files with product labels/tags using multiple DOCX libraries,
and provides a Tkinter GUI for user interaction.
"""
import sys, os, platform, subprocess, re, datetime, math, traceback, logging
import concurrent.futures
from io import BytesIO
from copy import deepcopy
from xml.sax.saxutils import unescape
import tkinter as tkmod
from tkinter import filedialog, messagebox, ttk
import pandas as pd
import numpy as np
from functools import lru_cache
from pathlib import Path
import io          


# ------------------ Third-Party DOCX Imports ------------------
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Mm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docxcompose.composer import Composer
import docxcompose
import os, datetime

# Paths — adjust if needed:
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates", "inventorySlips.docx")
OUTPUT_DIR    = os.path.expanduser("~/Downloads")


LOG_PATH = os.path.join(
    os.path.expanduser("~"),
    "Downloads",
    "lineage_change_log.csv"
)


import os, docxcompose
docxcompose_templates = (os.path.join(os.path.dirname(docxcompose.__file__), "templates"), "docxcompose/templates")

logging.basicConfig(level=logging.DEBUG, format="%(asctime)s %(levelname)s: %(message)s")
logging.debug("Application starting...")

from concurrent.futures import ThreadPoolExecutor
_IO_POOL = ThreadPoolExecutor(max_workers=1)          # single worker for I/O

from concurrent.futures import ThreadPoolExecutor
executor = ThreadPoolExecutor(max_workers=4)

def _add_cat_value(series, value):
    """Add *value* to a Categorical column if it isn't present."""
    if hasattr(series.dtype, "categories") and value not in series.cat.categories:
        return series.cat.add_categories([value])
    return series

_UPDATING_FILTERS = False




# Global variable to track which canvas is active.
# Global variables
current_canvas = None
available_tags_container = None
selected_tags_container = None
product_state_vars = {}
undo_stack = []  # For undo functionality
placeholder_img = None
print_vendor_back_var = None
SCALE_FACTOR = 1.0
WORD_WEIGHT = 5





import os, platform

def save_docx_and_replace(doc, path):
    """
    Save a python‑docx Document to `path`. If Word has that exact file open,
    it will close it for you (no changes saved), then overwrite it.
    """
    try:
        doc.save(path)
    except PermissionError:
        system = platform.system()
        if system == "Windows":
            from win32com.client import Dispatch
            word = Dispatch("Word.Application")
            for d in word.Documents:
                if os.path.abspath(d.FullName).lower() == os.path.abspath(path).lower():
                    d.Close(False)
                    break
        elif system == "Darwin":
            name = os.path.basename(path)
            applescript = f'''
            tell application "Microsoft Word"
              close (every document whose name is "{name}") saving no
            end tell
            '''
            os.system(f"osascript -e '{applescript}'")
        else:
            raise
        doc.save(path)


posabit_instructions = (
    "How to Obtain and Download Your Excel File from POSaBit\n\n"
    "1. Navigate to the POSaBit Inventory Page\n"
    "   Open your preferred web browser and log into your POSaBit account. Once logged in, navigate to the POSaBit → Inventory → Lots section. "
    "This is where you will be able to view all available inventory lots.\n\n"
    
    "2. Set Up Your Filters\n"
    "   On the left-hand side of the screen, you will see a filter sidebar. You need to apply the following filters to display only the relevant lots:\n"
    "       • Status: Change the status filter to \"Active\" so that only active inventory items are shown.\n"
    "       • Quantity On Hand: Adjust the filter to show only items with a Quantity On Hand above 0. "
    "This ensures you are only downloading items that are currently in stock.\n\n"
    
    "3. Run the Search\n"
    "   Once you have set the filters, click the \"Search\" button. This action will refresh the list of lots to display only those that match your filter criteria (active items with available quantity).\n\n"
    
    "4. Download Your Excel File\n"
    "   After your search results have been updated, locate the blue Download arrow button and click it. "
    "This will download an Excel file containing your filtered data.\n\n"
    
    "5. Upload the Excel File and Select a Template\n"
    "   Return to this application and use the Upload button provided to select and upload the Excel file you just downloaded. "
    "Once uploaded, choose the appropriate template for generating product labels or inventory slips."
)


# ------------------ Helper Functions ------------------
def set_current_canvas(event, canvas):
    global current_canvas
    current_canvas = canvas

def clear_current_canvas(event):
    global current_canvas
    current_canvas = None

def wrap_with_marker(text, marker):
    safe_text = str(text).replace('&', '&amp;')
    return f"{marker.upper()}_START{safe_text}{marker.upper()}_END"

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS  # set by PyInstaller
    except Exception:
        base_path = os.path.abspath(".")
    full_path = os.path.join(base_path, relative_path)
    print(f"Loading resource from: {full_path}")
    return full_path

def open_file(file_path):
    if not os.path.exists(file_path):
        logging.error("File not found: %s", file_path)
        return
    try:
        if platform.system() == "Windows":
            os.startfile(file_path)
        elif platform.system() == "Darwin":
            subprocess.Popen(["/usr/bin/open", "-a", "Microsoft Word", file_path])
        else:
            subprocess.Popen(["xdg-open", file_path])
    except Exception as e:
        logging.error("Error opening file: %s", e)


def get_default_upload_file():
    """
    Looks for files in the Downloads folder that start with "A Greener Today" and end with ".xlsx".
    Returns the full path of the most recently modified file, or None if no matching file is found.
    """
    downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
    files = []
    
    # Scan the Downloads directory for matching files.
    for f in os.listdir(downloads_dir):
        if f.startswith("A Greener Today") and f.lower().endswith(".xlsx"):
            full_path = os.path.join(downloads_dir, f)
            files.append(full_path)
    
    if files:
        # Use the most recently modified file.
        latest_file = max(files, key=os.path.getmtime)
        return latest_file
    else:
        return None

# ------------------ Global Variables and Font Schemes ------------------
global_df = None  # DataFrame from Excel file
product_check_vars = {}  # (Legacy: not used for filtering labels anymore)
selected_tags_vars = {}   # Dictionary for items moved to Selected Tag List (key: product name)
available_tags_vars = {}    # Dictionary for available tag list
move_history = []


FONT_SCHEME_HORIZONTAL = {
    "DESC": {"base_size": 28, "min_size": 12, "max_length": 100},
    "PRIC": {"base_size": 38, "min_size": 20, "max_length": 20},
    "LINEAGE": {"base_size": 20, "min_size": 12, "max_length": 30},
    "LINEAGE_CENTER": {"base_size": 18, "min_size": 12, "max_length": 30},
    "THC_CBD": {"base_size": 12, "min_size": 10, "max_length": 50},
    "RATIO": {"base_size": 10, "min_size": 8, "max_length": 30},
    "WEIGHT": {"base_size": 18, "min_size": 10, "max_length": 20},
    "UNITS": {"base_size": 18, "min_size": 10, "max_length": 20},
    "PRODUCTSTRAIN": {"base_size": 1, "min_size": 1, "max_length": 40},
    "PRODUCTBRAND_CENTER": {"base_size": 20, "min_size": 12, "max_length": 40}
}

FONT_SCHEME_VERTICAL = {
    "DESC": {"base_size": 23, "min_size": 12, "max_length": 100},
    "PRIC": {"base_size": 36, "min_size": 20, "max_length": 20},
    "LINEAGE": {"base_size": 18, "min_size": 12, "max_length": 30},
    "LINEAGE_CENTER": {"base_size": 18, "min_size": 12, "max_length": 30},
    "THC_CBD": {"base_size": 12, "min_size": 10, "max_length": 50},
    "RATIO": {"base_size": 8, "min_size": 10, "max_length": 30},
    "WEIGHT": {"base_size": 16, "min_size": 10, "max_length": 20},
    "UNITS": {"base_size": 16, "min_size": 10, "max_length": 20},
    "PRODUCTSTRAIN": {"base_size": 1, "min_size": 1, "max_length": 40},
    "PRODUCTBRAND_CENTER": {"base_size": 20, "min_size": 12, "max_length": 40}
}

FONT_SCHEME_MINI = {
    "DESC": {"base_size": 20, "min_size": 8, "max_length": 100},
    "PRIC": {"base_size": 22, "min_size": 10, "max_length": 20},
    "LINEAGE": {"base_size": 10, "min_size": 8, "max_length": 30},
    "LINEAGE_CENTER": {"base_size": 10, "min_size": 8, "max_length": 30},
    "THC_CBD": {"base_size": 8, "min_size": 6, "max_length": 50},
    "RATIO": {"base_size": 8, "min_size": 6, "max_length": 30},
    "WEIGHT": {"base_size": 10, "min_size": 8, "max_length": 20},
    "UNITS": {"base_size": 10, "min_size": 8, "max_length": 20},
    "PRODUCTSTRAIN": {"base_size": 1, "min_size": 1, "max_length": 40},
    "PRODUCTBRAND_CENTER": {"base_size": 7, "min_size": 1, "max_length": 40}
}

FONT_SCHEME_INVENTORY = {
    "DESC": {"base_size": 20, "min_size": 18, "max_length": 80},
    "PRIC": {"base_size": 30, "min_size": 24, "max_length": 15},
    "LINEAGE": {"base_size": 18, "min_size": 14, "max_length": 25},
    "LINEAGE_CENTER": {"base_size": 18, "min_size": 14, "max_length": 25},
    "THC_CBD": {"base_size": 14, "min_size": 12, "max_length": 50},
    "RATIO": {"base_size": 16, "min_size": 12, "max_length": 30},
    "WEIGHT": {"base_size": 20, "min_size": 18, "max_length": 15},
    "UNITS": {"base_size": 20, "min_size": 16, "max_length": 15},
    "PRODUCTSTRAIN": {"base_size": 18, "min_size": 12, "max_length": 40},
    "PRODUCTBRAND_CENTER": {"base_size": 22, "min_size": 16, "max_length": 40}
}

# ------------------ Helper Functions for Normalization ------------------
def normalize(val):
    return str(val).strip().lower()

def extract_float(x):
    try:
        matches = re.findall(r"[\d.]+", x)
        if matches:
            return float(matches[0])
    except Exception:
        pass
    return 0

# ------------------ UI Functions for Editing Data ------------------
def edit_data_manually():
    # Implementation for editing data manually
    file_path_val = file_entry.get()
    if not file_path_val:
        messagebox.showerror("Error", "Please upload a data file before editing.")
        return
    try:
        transformed_excel_file = preprocess_excel(file_path_val)
        logging.debug(f"Transformed file created: {transformed_excel_file}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to transform Excel file: {e}")
        return
    open_file(transformed_excel_file)
    response = messagebox.askokcancel(
        "Edit Data Manually",
        "The transformed spreadsheet has been opened in Excel.\n\n"
        "Please edit and save the file in Excel, then click OK to reload the updated data.\n"
        "If you haven't finished editing, click Cancel."
    )
    if response:
        try:
            global global_df
            global_df = pd.read_excel(transformed_excel_file, engine="openpyxl")
            populate_filter_dropdowns()
            populate_product_names()
            messagebox.showinfo("Reload Successful", "Data has been reloaded from the edited file.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to reload edited file: {e}")

# ------------------ DOCX Helper Functions ------------------
def disable_autofit(table):
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

def set_table_cell_spacing(table, spacing_twips):
    tblPr = table._element
    tblPr_obj = tblPr.find(qn('w:tblPr'))
    if tblPr_obj is None:
        tblPr_obj = OxmlElement('w:tblPr')
        tblPr.insert(0, tblPr_obj)
    tblCellSpacing = tblPr_obj.find(qn('w:tblCellSpacing'))
    if tblCellSpacing is None:
        tblCellSpacing = OxmlElement('w:tblCellSpacing')
        tblPr_obj.append(tblCellSpacing)
    tblCellSpacing.set(qn('w:w'), str(spacing_twips))
    tblCellSpacing.set(qn('w:type'), 'dxa')

def _set_row_height_exact(row, height_pt):
    trPr = row._tr.get_or_add_trPr()
    for child in trPr.findall(qn('w:trHeight')):
        trPr.remove(child)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_pt.pt * 20)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def cell_has_text(cell):
    return bool(cell.text.strip())

def rebuild_table_with_nonempty_cells(doc, old_table, num_cols=5):
    non_empty_texts = [
    cell.text for row in old_table.rows for cell in row.cells if cell_has_text(cell)
]

    # Create new table directly using texts:
    new_table = doc.add_table(rows=num_rows, cols=num_cols)
    idx = 0
    for row in new_table.rows:
        for cell in row.cells:
            if idx < len(non_empty_texts):
                cell.text = non_empty_texts[idx]
                idx += 1
    old_table._element.getparent().remove(old_table._element)
    total_cells = len(non_empty_cells)
    if total_cells == 0:
        return None
    rows_needed = math.ceil(total_cells / num_cols)
    new_table = doc.add_table(rows=rows_needed, cols=num_cols)
    new_table.alignment = 1
    disable_autofit(new_table)
    tblPr = new_table._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    new_table._element.insert(0, tblPr)
    tblGrid = OxmlElement('w:tblGrid')
    fixed_col_width = "2000"
    for _ in range(num_cols):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), fixed_col_width)
        tblGrid.append(gridCol)
    new_table._element.insert(0, tblGrid)
    idx = 0
    for r in range(rows_needed):
        for c in range(num_cols):
            cell = new_table.cell(r, c)
            cell._tc.clear_content()
            if idx < total_cells:
                for child in non_empty_cells[idx]:
                    cell._tc.append(deepcopy(child))
                idx += 1
            else:
                cell.text = ""
    return new_table

# ------------------ DOCX Expand Template Functions ------------------


def expand_template_to_3x3_fixed(template_path):
    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("Template must contain at least one table.")
    old_table = doc.tables[0]
    source_cell_xml = deepcopy(old_table.cell(0, 0)._tc)
    old_table._element.getparent().remove(old_table._element)
    # strip any leading empty paragraphs
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    new_table = doc.add_table(rows=3, cols=3)
    new_table.alignment = 1
    disable_autofit(new_table)

    # ── ADD ONLY INTERIOR CUT-GUIDELINES ──
    tblPr = new_table._element.find(qn('w:tblPr'))
    # remove any existing borders
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)

    tblBorders = OxmlElement('w:tblBorders')
    # hide outer borders
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn('w:val'), "nil")
        tblBorders.append(bd)
    # draw interior lines
    for side in ("insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn('w:val'), "single")
        bd.set(qn('w:sz'), "4")
        bd.set(qn('w:color'), "D3D3D3")
        bd.set(qn('w:space'), "0")
        tblBorders.append(bd)
    tblPr.append(tblBorders)

    # rebuild grid to fixed 3.5"x2.5" cells
    fixed_col_width = str(int(3.5 * 1440 / 3))  # total width split among 3
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(3):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), fixed_col_width)
        tblGrid.append(gridCol)
    new_table._element.insert(0, tblGrid)
    for i in range(3):
        for j in range(3):
            label_num = i * 3 + j + 1
            cell = new_table.cell(i, j)
            cell._tc.clear_content()
            new_tc = deepcopy(source_cell_xml)
            for text_el in new_tc.iter():
                if text_el.tag == qn('w:t') and text_el.text and "Label1" in text_el.text:
                    text_el.text = text_el.text.replace("Label1", f"Label{label_num}")
            cell._tc.extend(new_tc.xpath("./*"))
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def expand_template_to_5x6_fixed_scaled(template_path, num_inputs, scale_factor=1.0):
    from docx import Document
    num_cols = 5
    num_rows = math.ceil(num_inputs / num_cols)
    width_twips = str(int(3.5 * 1440 / num_cols))
    height_points = Pt(2.5 * 72 / (9/num_rows))  # keep roughly 2.5" tall per row if 3 rows; adjust as needed

    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("Template must contain at least one table.")
    old_table = doc.tables[0]
    source_cell_xml = deepcopy(old_table.cell(0, 0)._tc)
    old_table._element.getparent().remove(old_table._element)
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    new_table = doc.add_table(rows=num_rows, cols=num_cols)
    new_table.alignment = 1
    disable_autofit(new_table)

    # ── ADD ONLY INTERIOR CUT-GUIDELINES ──
    tblPr = new_table._element.find(qn('w:tblPr'))
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)

    tblBorders = OxmlElement('w:tblBorders')
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn('w:val'), "nil")
        tblBorders.append(bd)
    for side in ("insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn('w:val'), "single")
        bd.set(qn('w:sz'), "4")
        bd.set(qn('w:color'), "D3D3D3")
        bd.set(qn('w:space'), "0")
        tblBorders.append(bd)
    tblPr.append(tblBorders)

    # fixed grid columns
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), width_twips)
        tblGrid.append(gridCol)
    new_table._element.insert(0, tblGrid)
    # set exact row heights
    for row in new_table.rows:
        row.height = height_points
        _set_row_height_exact(row, height_points)

    # populate cells
    label_num = 1
    for i in range(num_rows):
        for j in range(num_cols):
            cell = new_table.cell(i, j)
            new_tc = deepcopy(source_cell_xml)
            for t in new_tc.iter(qn('w:t')):
                if t.text and "Label1" in t.text:
                    t.text = t.text.replace("Label1", f"Label{label_num}")
            cell._tc.clear_content()
            if label_num <= num_inputs:
                cell._tc.extend(new_tc.xpath("./*"))
            label_num += 1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def expand_template_to_2x2_inventory_slips(template_path):
    num_rows, num_cols = 2, 2
    width_twips = str(int(3.5 * 1440 / num_cols))
    height_points = Pt(2.5 * 72 / num_rows)

    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("Template must contain at least one table.")
    old_table = doc.tables[0]
    source_cell_xml = deepcopy(old_table.cell(0, 0)._tc)
    old_table._element.getparent().remove(old_table._element)
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    new_table = doc.add_table(rows=num_rows, cols=num_cols)
    new_table.alignment = 1
    disable_autofit(new_table)

    # ── ADD ONLY INTERIOR CUT-GUIDELINES ──
    tblPr = new_table._element.find(qn('w:tblPr'))
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)

    tblBorders = OxmlElement('w:tblBorders')
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn('w:val'), "nil")
        tblBorders.append(bd)
    for side in ("insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn('w:val'), "single")
        bd.set(qn('w:sz'), "4")
        bd.set(qn('w:color'), "D3D3D3")
        bd.set(qn('w:space'), "0")
        tblBorders.append(bd)
    tblPr.append(tblBorders)

    # fixed grid columns & rows
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), width_twips)
        tblGrid.append(gridCol)
    new_table._element.insert(0, tblGrid)
    for row in new_table.rows:
        row.height = height_points
        _set_row_height_exact(row, height_points)

    # populate with inventory slip placeholders
    for i in range(num_rows):
        for j in range(num_cols):
            label_num = i * num_cols + j + 1
            cell = new_table.cell(i, j)
            cell._tc.clear_content()
            new_tc = deepcopy(source_cell_xml)
            for text_el in new_tc.iter():
                if text_el.tag == qn('w:t') and text_el.text and "Label1" in text_el.text:
                    text_el.text = text_el.text.replace("Label1", f"Label{label_num}")
            cell._tc.extend(new_tc.xpath("./*"))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer



# ------------------ Autosize and Conditional Formatting ------------------
def set_run_font_size(run, font_size):
    run.font.size = font_size
    sz_val = str(int(font_size.pt * 2))
    rPr = run._element.get_or_add_rPr()
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), sz_val)

from docx.shared import Pt

def _complexity(text):
    """Combine character count and weighted word count into one score."""
    text = str(text or "")
    return len(text) + len(text.split()) * WORD_WEIGHT


def get_thresholded_font_size_by_word_count(text, orientation='vertical', scale_factor=1.0):
    comp = _complexity(text)
    o = orientation.lower()

    if o == 'mini':
        # e.g. DESC on mini tags
        if comp <  30: size = 18
        elif comp <  40: size = 17
        elif comp <  50: size = 16
        elif comp < 75: size = 14
        elif comp <  100: size = 12
        elif comp < 120: size = 11
        else:           size = 9

    elif o == 'vertical':
        if comp <  30: size = 29
        elif comp <  60: size = 26
        elif comp < 100: size = 22
        elif comp < 140: size = 20
        else:           size = 18

    elif o == 'horizontal':
        if comp <  20: size = 34
        elif comp <  30: size = 32
        elif comp <  50: size = 28
        elif comp <  60: size = 26
        elif comp < 100: size = 24
        elif comp < 120: size = 22
        else:           size = 20

    elif o == 'inventory':
        if comp <  30: size = 18
        elif comp <  60: size = 16
        elif comp <  90: size = 14
        elif comp < 120: size = 12
        elif comp < 180: size = 10
        else:           size = 8

    else:  # fallback
        size = 14

    return Pt(size * scale_factor)


def get_thresholded_font_size_ratio(text, orientation='vertical', scale_factor=1.0):
    comp = _complexity(text)
    o = orientation.lower()

    if o == 'mini':
        if comp <  20: size =  8
        elif comp <  40: size =  7
        else:           size =  6

    elif o == 'vertical':
        if comp <  20: size = 14
        elif comp <  30: size = 12
        elif comp < 100: size = 10
        elif comp < 140: size =  8
        else:           size = 10

    elif o == 'horizontal':
        if comp <  20: size = 16
        elif comp <  30: size = 14
        elif comp <  0: size = 12
        elif comp < 100: size = 10
        else:           size = 10

    elif o == 'inventory':
        if comp <  30: size = 18
        elif comp <  60: size = 16
        elif comp <  90: size = 14
        elif comp < 120: size = 12
        elif comp < 180: size = 10
        else:           size =  8

    else:
        size = 10

    return Pt(size * scale_factor)


def get_thresholded_font_size_brand(text, orientation='vertical', scale_factor=1.0):
    comp = _complexity(text)
    o = orientation.lower()

    if o == 'mini':
        if comp <  10: size = 12
        elif comp <  20: size = 10
        elif comp <  50: size =  8
        else:           size =  7

    elif o == 'vertical':
        if comp <  20: size = 16
        elif comp <  40: size = 14
        elif comp <  80: size = 12
        else:           size = 11

    elif o == 'horizontal':
        if comp <  20: size = 18
        elif comp <  40: size = 16
        elif comp <  80: size = 12
        else:           size = 10

    elif o == 'inventory':
        if comp <  30: size = 18
        elif comp <  60: size = 16
        elif comp <  90: size = 14
        elif comp < 120: size = 12
        elif comp < 180: size = 10
        else:           size =  8

    else:
        size = 14

    return Pt(size * scale_factor)


def autosize_field_in_paragraph(para, marker_start, marker_end, font_params, orientation, font_name="Arial", bold=True, scale_factor=1.0):
    full_text = "".join(run.text for run in para.runs)
    if marker_start in full_text and marker_end in full_text:
        try:
            field_text = full_text.split(marker_start)[1].split(marker_end)[0].strip()
        except IndexError:
            return
        # Debug print:
        print(f"[DEBUG] Field text for marker {marker_start}: '{field_text}'")
        if marker_start == "PRODUCTBRAND_CENTER_START":
            new_size_val = get_thresholded_font_size_brand(field_text, orientation, scale_factor)
        elif marker_start == "DESC_START":
            new_size_val = get_thresholded_font_size_by_word_count(field_text, orientation, scale_factor)
        elif marker_start == "RATIO_START":
            new_size_val = get_thresholded_font_size_ratio(field_text, orientation, scale_factor)
        elif marker_start == "PRODUCTSTRAIN_START":
            new_size_val = Pt(font_params["base_size"])
        else:
            length = len(field_text)
            base_size = font_params["base_size"]
            max_length = font_params["max_length"]
            min_size = font_params["min_size"]
            new_size_val = Pt(base_size) if length <= max_length else Pt(max(min_size, base_size * (max_length / length)))
        # Increase size for placeholders.
        bold = True

        new_text = unescape(full_text.replace(marker_start, "").replace(marker_end, ""))
        p_element = para._element
        for child in list(p_element):
            p_element.remove(child)
        new_run = para.add_run(new_text)
        new_run.font.size = new_size_val
        new_run.font.name = font_name
        new_run.font.bold = bold
        set_run_font_size(new_run, new_size_val)
        if marker_start == "PRODUCTBRAND_CENTER_START":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def autosize_fields(doc, font_scheme, orientation, scale_factor=1.0):
    def recursive_autosize(element, marker_start, marker_end, font_params, orientation, scale_factor):
        for para in element.paragraphs:
            autosize_field_in_paragraph(para, marker_start, marker_end, font_params, orientation, scale_factor=scale_factor)
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    recursive_autosize(cell, marker_start, marker_end, font_params, orientation, scale_factor)
    recursive_autosize(doc, "DESC_START", "DESC_END", font_scheme["DESC"], orientation, scale_factor)
    recursive_autosize(doc, "PRIC_START", "PRIC_END", font_scheme["PRIC"], orientation, scale_factor)
    recursive_autosize(doc, "LINEAGE_START", "LINEAGE_END", font_scheme["LINEAGE"], orientation, scale_factor)
    recursive_autosize(doc, "LINEAGE_CENTER_START", "LINEAGE_CENTER_END", font_scheme["LINEAGE_CENTER"], orientation, scale_factor)
    recursive_autosize(doc, "THC_CBD_START", "THC_CBD_END", font_scheme["THC_CBD"], orientation, scale_factor)
    recursive_autosize(doc, "RATIO_START", "RATIO_END", font_scheme["DESC"], orientation, scale_factor)
    recursive_autosize(doc, "WEIGHT_START", "WEIGHT_END", font_scheme["WEIGHT"], orientation, scale_factor)
    recursive_autosize(doc, "UNITS_START", "UNITS_END", font_scheme["UNITS"], orientation, scale_factor)
    recursive_autosize(doc, "PRODUCTSTRAIN_START", "PRODUCTSTRAIN_END", font_scheme["PRODUCTSTRAIN"], orientation, scale_factor)
    recursive_autosize(doc, "PRODUCTBRAND_CENTER_START", "PRODUCTBRAND_CENTER_END", font_scheme["PRODUCTBRAND_CENTER"], orientation, scale_factor)
    return doc

def apply_conditional_formatting(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                apply_formatting_to_cell(cell)
    shade_middle_row_if_needed(doc)
    return doc

def apply_formatting_to_cell(cell):
    text = cell.text.strip().upper()

    # 1) CANNABINOID RATIOS → SHADE YELLOW
    if any(chem in text for chem in ["CBD", "CBN", "CBG", "CBC"]):
        set_cell_background(cell, "F1C232")
        set_font_color_white(cell)
        return

    # 2) PARAPHERNALIA → SHADE PINK
    if "PARAPHERNALIA" in text:
        set_cell_background(cell, "FFC0CB")
        set_font_color_white(cell)
        return

    # 3) HYBRID SUB‑TYPES
    if "HYBRID/INDICA" in text or "HYBRID INDICA" in text:
        set_cell_background(cell, "9900FF"); set_font_color_white(cell); return
    if "HYBRID/SATIVA" in text or "HYBRID SATIVA" in text:
        set_cell_background(cell, "ED4123"); set_font_color_white(cell); return

    # 4) PLAIN LINEAGES
    if "INDICA"   in text:
        set_cell_background(cell, "9900FF"); set_font_color_white(cell); return
    if "SATIVA"   in text:
        set_cell_background(cell, "ED4123"); set_font_color_white(cell); return
    if "HYBRID"   in text:
        set_cell_background(cell, "009900"); set_font_color_white(cell); return
    if "MIXED"    in text:
        set_cell_background(cell, "0021F5"); set_font_color_white(cell); return
    if "CBD"      in text:  # fallback for lone “CBD”
        set_cell_background(cell, "F1C232"); set_font_color_white(cell); return
    if "PARAPHERNALIA" in text:  # fallback catch
        set_cell_background(cell, "FFC0CB"); set_font_color_white(cell); return

    # 5) DEFAULT → white background
    set_cell_background(cell, "FFFFFF")





def set_cell_background(cell, color_hex):
    if not cell.text.strip():
        cell.text = " "
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    new_shd = OxmlElement('w:shd')
    new_shd.set(qn('w:val'), 'clear')
    new_shd.set(qn('w:color'), 'auto')
    new_shd.set(qn('w:fill'), color_hex.upper())
    tcPr.append(new_shd)

def set_font_color_white(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.highlight_color = None

def shade_middle_row_if_needed(doc, shade_color="C0C0C0"):
    for table in doc.tables:
        if not table.rows:
            continue
        mid_index = len(table.rows) // 2
        mid_row = table.rows[mid_index]
        row_text = " ".join(cell.text for cell in mid_row.cells).upper()
        if "LINEAGE" in row_text or "PRODUCT STRAIN" in row_text:
            for cell in mid_row.cells:
                set_cell_background(cell, shade_color)
                set_font_color_white(cell)
    return doc

def safe_fix_paragraph_spacing(doc):
    for para in doc.paragraphs:
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
    return doc

def remove_extra_spacing(doc):
    try:
        normal_style = doc.styles["Normal"].paragraph_format
        normal_style.space_before = Pt(0)
        normal_style.space_after = Pt(0)
        normal_style.line_spacing = 1
    except Exception as e:
        print("Error adjusting Normal style:", e)
    return doc

def clear_cell_margins(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                for margin in ("top", "left", "bottom", "right"):
                    m = tcPr.find(qn(f"w:{margin}"))
                    if m is None:
                        m = OxmlElement(f"w:{margin}")
                        tcPr.append(m)
                    m.set(qn("w:w"), "0")
                    m.set(qn("w:type"), "dxa")
    return doc

def clear_table_cell_padding(doc):
    for table in doc.tables:
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is not None:
            tblCellMar = tblPr.find(qn('w:tblCellMar'))
            if tblCellMar is not None:
                for margin in ['top', 'left', 'bottom', 'right']:
                    m = tblCellMar.find(qn(f'w:{margin}'))
                    if m is not None:
                        m.set(qn('w:w'), "0")
                        m.set(qn('w:type'), "dxa")
    return doc

def compact_table_cells(doc, num_cols=3):
    if not doc.tables:
        return doc
    orig_table = doc.tables[0]
    non_blank_cells = []
    for row in orig_table.rows:
        for cell in row.cells:
            if cell.text.strip():
                non_blank_cells.append(deepcopy(cell._tc))
    orig_table._element.getparent().remove(orig_table._element)
    num_cells = len(non_blank_cells)
    num_rows = (num_cells + num_cols - 1) // num_cols
    new_table = doc.add_table(rows=num_rows, cols=num_cols)
    new_table.alignment = 1
    fixed_col_width = "2000"
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), fixed_col_width)
        tblGrid.append(gridCol)
    new_table._element.insert(0, tblGrid)
    cell_index = 0
    for r in range(num_rows):
        for c in range(num_cols):
            cell = new_table.cell(r, c)
            cell._tc.clear_content()
            if cell_index < num_cells:
                for child in non_blank_cells[cell_index]:
                    cell._tc.append(deepcopy(child))
                cell_index += 1
            else:
                cell.text = ""
    return new_table

def reapply_table_cell_spacing_only(doc, spacing_inches=0.03):
    spacing_twips = int(spacing_inches * 1440)
    for table in doc.tables:
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._element.insert(0, tblPr)
        tblCellSpacing = tblPr.find(qn('w:tblCellSpacing'))
        if tblCellSpacing is None:
            tblCellSpacing = OxmlElement('w:tblCellSpacing')
            tblPr.append(tblCellSpacing)
        tblCellSpacing.set(qn('w:w'), str(spacing_twips))
        tblCellSpacing.set(qn('w:type'), 'dxa')

def remove_trailing_blank_paragraphs(doc):
    """
    Remove trailing empty paragraphs from a Document to help prevent a blank page.
    """
    # Iterate over paragraphs in reverse order
    for para in reversed(doc.paragraphs):
        if not para.text.strip():  # if paragraph is blank
            # Remove the paragraph element from its parent
            p_element = para._element
            p_element.getparent().remove(p_element)
        else:
            # Stop once a non-empty paragraph is reached.
            break
    return doc

# ------------------ Excel Processing Functions ------------------
from decimal import Decimal, InvalidOperation

def format_price(p):
    try:
        value = str(p).strip().lstrip("$")
        val = float(value)
        if val.is_integer():
            return f"'{int(val)}"
        else:
            s = str(val).rstrip("0").rstrip(".")
            return f"'{s}"
    except Exception:
        return f"'{str(p).strip().lstrip('$')}"

    
def format_weight(w):
    try:
        val = float(w)
        return str(int(val)) if val.is_integer() else str(val)
    except Exception:
        return str(w)

def sanitize_filename(s):
    allowed = "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789 _-&"
    return "".join(ch for ch in s if ch in allowed).replace(" ", "_")


def format_price_preprocess(p):
    try:
        s = str(p).strip()
        # Remove any leading "$"
        if s.startswith("$"):
            s = s[1:]
        # Remove apostrophes and stray whitespace
        s = s.replace("'", "").strip()
        # Convert to float
        val = float(s)
        # If the value is integer, return it without any decimal portion
        if val.is_integer():
            return f"${int(val)}"
        else:
            # Otherwise, format with 2 decimals, then remove any trailing zeros and dot
            formatted = f"{val:.2f}"
            # If formatted ends with .00, remove it completely
            if formatted.endswith(".00"):
                return f"${formatted[:-3]}"
            else:
                formatted = formatted.rstrip("0").rstrip(".")
                return f"${formatted}"
    except Exception:
        # Fallback just in case
        fallback = str(p).strip()
        if fallback.startswith("$"):
            fallback = fallback[1:]
        fallback = fallback.replace("'", "").strip()
        try:
            fv = float(fallback)
            if fv.is_integer():
                return f"${int(fv)}"
            else:
                formatted = f"{fv:.2f}"
                if formatted.endswith(".00"):
                    return f"${formatted[:-3]}"
                else:
                    return f"${formatted.rstrip('0').rstrip('.')}"
        except Exception:
            return f"${fallback}"
        
def fix_description_spacing(desc):
    """
    Inserts a space before a hyphen that is immediately followed by a digit.
    
    For example:
      "Gelato Infused Pre-Roll- 1g" becomes "Gelato Infused Pre-Roll - 1g"
      
    This function only affects hyphens directly preceding digits,
    so it leaves compound words like "Pre-Roll" intact.
    """
    # (?<!\s) ensures that the hyphen isn't already preceded by a space.
    # \s* eats up any existing whitespace after the hyphen.
    # (\d) captures the first digit that follows.
    return re.sub(r'(?<!\s)-\s*(\d)', r' - \1', desc)

def preprocess_excel(file_path, filters=None):
    import datetime, os, re
    import numpy as np
    import pandas as pd

    # 1) Read & dedupe, force-key columns to string for .str ops
    dtype_dict = {
        "Product Type*": "string",
        "Lineage": "string",
        "Product Brand": "string",
        "Vendor": "string",
        "Weight Unit* (grams/gm or ounces/oz)": "string",
        "Product Name*": "string"
    }
    df = pd.read_excel(file_path, engine="openpyxl", dtype=dtype_dict)
    df.drop_duplicates(inplace=True)

    # 2) Trim product names
    if "Product Name*" in df.columns:
        df["Product Name*"] = df["Product Name*"].str.lstrip()

    # 3) Ensure required columns exist
    for col in ["Product Type*", "Lineage", "Product Brand"]:
        if col not in df.columns:
            df[col] = "Unknown"

    # 4) Exclude sample rows
    df = df[~df["Product Type*"].isin(["Samples - Educational", "Sample - Vendor"])]

    # 5) Rename for convenience
    df.rename(columns={
        "Weight Unit* (grams/gm or ounces/oz)": "Units",
        "Price* (Tier Name for Bulk)": "Price",
        "Vendor/Supplier*": "Vendor",
        "DOH Compliant (Yes/No)": "DOH",
        "Concentrate Type": "Ratio"
    }, inplace=True)

    # 6) Normalize units
    if "Units" in df.columns:
        df["Units"] = df["Units"].str.lower().replace(
            {"ounces": "oz", "grams": "g"}, regex=True
        )

    # 7) Standardize Lineage
    if "Lineage" in df.columns:
        df["Lineage"] = (
            df["Lineage"]
              .str.lower()
              .replace({
                  "indica_hybrid": "HYBRID/INDICA",
                  "sativa_hybrid": "HYBRID/SATIVA",
                  "sativa": "SATIVA",
                  "hybrid": "HYBRID",
                  "indica": "INDICA"
              })
              .fillna("HYBRID")
              .str.upper()
        )

    # 8) Build Description & Ratio & Strain
    if "Product Name*" in df.columns:
        df["Description"] = df["Product Name*"].str.split(" by").str[0]
        mask_para = df["Product Type*"].str.strip().str.lower() == "paraphernalia"
        df.loc[mask_para, "Description"] = (
            df.loc[mask_para, "Description"]
              .str.replace(r"\s*-\s*\d+g$", "", regex=True)
        )
           # … after this block that builds df["Ratio"] …
        df["Ratio"] = df["Product Name*"].str.extract(r"-\s*(.+)").fillna("")
        df["Ratio"] = df["Ratio"].str.replace(r" / ", " ", regex=True)

            # ── ensure “Product Strain” exists and is a Categorical ──────────────
        if "Product Strain" not in df.columns:
            df["Product Strain"] = ""
        # this will turn anything in that column into a category dtype
        df["Product Strain"] = df["Product Strain"].astype("category")

            # ‑‑‑ Force all non‑CBD‑Blend strains to "Mixed" ───────────────────────
        if "Product Strain" in df.columns:
            # Convert to plain string then override
            df["Product Strain"] = df["Product Strain"].astype(str).apply(
                lambda s: "CBD Blend" if s == "CBD Blend" else "Mixed"
            ).astype("category")


        # ── now force CBD Blend for any ratio containing CBD, CBC, CBN or CBG ──
        mask_cbd_ratio = df["Ratio"].str.contains(
            r"\b(?:CBD|CBC|CBN|CBG)\b", case=False, na=False
        )
        if mask_cbd_ratio.any():
            # add “CBD Blend” to the categories if it’s not already there
            if "CBD Blend" not in df["Product Strain"].cat.categories:
                df["Product Strain"] = df["Product Strain"].cat.add_categories(["CBD Blend"])
            # assign
            df.loc[mask_cbd_ratio, "Product Strain"] = "CBD Blend"



        # ── SPECIAL CASE: anything with Product Type “paraphernalia” gets
    #    its Product Strain forcibly set to "Paraphernalia"
    mask_para = df["Product Type*"].str.strip().str.lower() == "paraphernalia"

    # ensure the column exists as categorical and add the new category
    if "Product Strain" not in df.columns:
        df["Product Strain"] = pd.Categorical([], categories=["Paraphernalia"])
    else:
        # if it’s already categorical, just add the new category
        if isinstance(df["Product Strain"].dtype, pd.CategoricalDtype):
            if "Paraphernalia" not in df["Product Strain"].cat.categories:
                df["Product Strain"] = df["Product Strain"].cat.add_categories(["Paraphernalia"])
        else:
            # not categorical yet → make it categorical with this extra
            df["Product Strain"] = pd.Categorical(df["Product Strain"], 
                                                  categories=list(df["Product Strain"].unique()) + ["Paraphernalia"])

    # now you can safely assign
    df.loc[mask_para, "Product Strain"] = "Paraphernalia"


    # 9) Convert key fields to categorical
    for col in ["Product Type*", "Lineage", "Product Brand", "Vendor"]:
        if col in df.columns:
            df[col] = df[col].astype("category")

    # 10) CBD overrides
    if "Description" in df.columns and "Lineage" in df.columns:
        cbd_mask = df["Description"].str.contains(
            r"CBD|CBN|CBC|CBG|:", case=False, na=False
        )
        if "CBD" not in df["Lineage"].cat.categories:
            df["Lineage"] = df["Lineage"].cat.add_categories(["CBD"])
        df.loc[cbd_mask, "Lineage"] = "CBD"
    if "Description" in df.columns and "Product Strain" in df.columns:
        cbd_mask = df["Description"].str.contains(
            r"CBD|CBN|CBC|CBG|:", case=False, na=False
        )
        if "CBD Blend" not in df["Product Strain"].cat.categories:
            df["Product Strain"] = df["Product Strain"].cat.add_categories(["CBD Blend"])
        df.loc[cbd_mask, "Product Strain"] = "CBD Blend"

    # 11) Trim any extra columns
    if df.shape[1] > 41:
        df = df.iloc[:, :41]

    # 12) Normalize Weight* and CombinedWeight
    if "Weight*" in df.columns:
        df["Weight*"] = pd.to_numeric(df["Weight*"], errors="coerce") \
            .apply(lambda x: str(int(x)) if pd.notnull(x) and float(x).is_integer() else str(x))
    if "Weight*" in df.columns and "Units" in df.columns:
        df["CombinedWeight"] = (df["Weight*"] + df["Units"]).astype("category")

    # 13) Format Price
    if "Price" in df.columns:
        def format_p(p):
            s = str(p).strip().lstrip("$").replace("'", "").strip()
            try:
                v = float(s)
                return f"${int(v)}" if v.is_integer() else f"${v:.2f}"
            except:
                return f"${s}"
        df["Price"] = df["Price"].apply(lambda x: format_p(x) if pd.notnull(x) else "")
        df["Price"] = df["Price"].astype("string")

    # 14) Special pre-roll Ratio logic
    def process_ratio(row):
        t = str(row.get("Product Type*", "")).strip().lower()
        if t in ["pre-roll", "infused pre-roll"]:
            parts = str(row.get("Ratio", "")).split(" - ")
            if len(parts) >= 3:
                new = " - ".join(parts[2:]).strip()
            elif len(parts) == 2:
                new = parts[1].strip()
            else:
                new = parts[0].strip()
            return f" - {new}" if not new.startswith(" - ") else new
        return row.get("Ratio", "")
    df["Ratio"] = df.apply(process_ratio, axis=1)

    # … (suffix-building and Excel output) …
    today = datetime.datetime.today().strftime("%Y-%m-%d")
    suffix = "all"  # or built from `filters`
    out = os.path.join(
        os.path.expanduser("~"),
        "Downloads",
        f"{today}_{suffix}.xlsx"
    )
    df.to_excel(out, index=False, engine="openpyxl")
    return out


def chunk_records(records, chunk_size=4):
    """Yield successive n‑sized chunks from the list of records."""
    for i in range(0, len(records), chunk_size):
        yield records[i:i + chunk_size]

def no_filters_selected():
    filters = [
        product_type_filter_var.get(),
        lineage_filter_var.get(),
        product_brand_filter_var.get(),
        vendor_filter_var.get(),
        weight_filter_var.get(),
        product_strain_filter_var.get()
    ]
    return all(f == "All" for f in filters)



# ------------------ Processing Functions ------------------
def process_chunk(args):
    """
    Processes a chunk of records and returns a DOCX document as a BytesIO buffer.
    
    For the "inventory" orientation, the function assumes 4 cells per slip.
    Extra fields for inventory (AcceptedDate, Vendor, Barcode, ProductName, ProductType, QuantityReceived)
    are added to each label context.
    """
    from docx.shared import Mm
    from io import BytesIO
    chunk, base_template, font_scheme, orientation, fixed_scale = args
    if orientation == "mini":
        local_template_buffer = expand_template_to_5x6_fixed_scaled(base_template, num_inputs=len(chunk), scale_factor=fixed_scale)
        tpl = DocxTemplate(local_template_buffer)
    else:
        tpl = DocxTemplate(base_template)
    
    context = {}
    image_width = Mm(8) if orientation == "mini" else Mm(12 if orientation == 'vertical' else 14)
    doh_image_path = resource_path(os.path.join("templates", "DOH.png"))
    
    if orientation == "mini":
        num_labels = 30
    elif orientation == "inventory":
        num_labels = 4
    else:
        num_labels = 9

    for i in range(num_labels):
        label_data = {}
        if i < len(chunk):
            row = chunk[i]
            doh_value = str(row.get("DOH", "")).strip()
            product_type = str(row.get("Product Type*", "")).strip().lower()
            if doh_value == "Yes":
                high_cbd_types = [
                    "high cbd edible liquid - doh compliant",
                    "high cbd edible solid - doh compliant",
                    "high cbd topical - doh compliant"
                ]
                if product_type in high_cbd_types:
                    high_cbd_image_path = resource_path(os.path.join("templates", "HighCBD.png"))
                    label_data["DOH"] = InlineImage(tpl, high_cbd_image_path, width=image_width)
                else:
                    label_data["DOH"] = InlineImage(tpl, doh_image_path, width=image_width)
            else:
                label_data["DOH"] = ""
                
            price_val = f"{row.get('Price', '')}"
            label_data["Price"] = wrap_with_marker(price_val, "PRIC")
            
            lineage_text   = str(row.get("Lineage", "")).strip()
            product_brand  = str(row.get("Product Brand", "")).strip()
            label_data["ProductBrand"] = wrap_with_marker(product_brand.upper(), "PRODUCTBRAND_CENTER")

            # ── SPECIAL CASE: paraphernalia shows Vendor instead of Brand ──
          # ── SPECIAL CASE: paraphernalia ──────────────────────────────────
            if orientation not in ["mini", "inventory"] and product_type == "paraphernalia":
                vendor_text = str(row.get("Vendor", "")).strip()
                # show vendor in the Lineage cell
                label_data["Lineage"]         = wrap_with_marker(vendor_text.upper(), "PRODUCTBRAND_CENTER")
                # no THC/CBD block
                label_data["Ratio_or_THC_CBD"] = ""
                # force Product Strain to read "Paraphernalia"
                label_data["ProductStrain"]    = wrap_with_marker("Paraphernalia", "PRODUCTSTRAIN")
                # remove any weight/units field
                label_data["WeightUnits"]      = ""


            # ── all other types unchanged ─────────────────────────────────
            elif orientation not in ["mini", "inventory"]:
                classic_types = {
                    "flower", "vape cartridge", "solventless concentrate",
                    "concentrate", "pre-roll", "infused pre-roll"
                }
                if product_type in classic_types:
                    label_data["Lineage"]         = wrap_with_marker(lineage_text, "LINEAGE")
                    label_data["Ratio_or_THC_CBD"] = wrap_with_marker("THC:\n\nCBD:", "THC_CBD")
                    label_data["ProductStrain"]    = ""
                else:
                    label_data["Lineage"]         = wrap_with_marker(product_brand.upper(), "PRODUCTBRAND_CENTER")
                    label_data["Ratio_or_THC_CBD"] = wrap_with_marker(row.get("Ratio", ""), "RATIO")
                    label_data["ProductStrain"]    = wrap_with_marker(row.get("Product Strain", ""), "PRODUCTSTRAIN")
            else:
                # mini/inventory
                label_data["Lineage"]         = ""
                label_data["Ratio_or_THC_CBD"] = ""
                label_data["ProductStrain"]    = ""

            
            label_data["ProductBrandFontSize"] = get_thresholded_font_size_brand(product_brand, scale_factor=1.0)
            
            def format_ratio_multiline(ratio_text):
                if not isinstance(ratio_text, str):
                    return ""
                parts = re.split(r"\s*\|\s*|\s{2,}", ratio_text.strip())
                return "\n".join(p.strip() for p in parts if p.strip())
            import re   # at top of your MAIN.py

            # … after you’ve pulled `product_type = str(row.get("Product Type*", "")).lower()` …

            label_data["Ratio"] = wrap_with_marker(
                format_ratio_multiline(row.get("Ratio", "")),
                "RATIO"
            )

            # clean up the description cell:
            raw_desc = str(row.get("Description", "")).strip()
            if product_type == "paraphernalia":
                # drop any trailing " - Xg" (e.g. " - 0g")
                cleaned_desc = re.sub(r"\s*-\s*\d+g$", "", raw_desc)
            else:
                cleaned_desc = raw_desc

                

            label_data["Description"] = wrap_with_marker(cleaned_desc, "DESC")

            
                       # … after you’ve extracted raw_desc …
            # get the numeric weight and its unit
            try:
                weight_val = float(row.get("Weight*", ""))
            except Exception:
                weight_val = None
            units_val = row.get("Units", "").lower()    # e.g. 'g' or 'oz'

            # ── NEW: convert certain gram‑based products to oz ─────────────
            edible_types = {
                "edible (solid)",
                "edible (liquid)",
                "high cbd edible liquid",   # treat this the same as “edible (liquid)”
                "tincture",
                "topical",
                "capsule",
}

            if product_type in edible_types and units_val in {"g", "grams"} and weight_val is not None:
                weight_val = weight_val * 0.03527396195
                units_val = "oz"

            # now build the display string
            if weight_val is not None and units_val:
                weight_str = f"{weight_val:.2f}".rstrip("0").rstrip(".")
                weight_units = f" -\u00A0{weight_str}{units_val}"
            else:
                weight_units = ""

            # … earlier in process_chunk, after you have:
            # … after reading weight_val, units_val …
            product_type = product_type.lower()
            units_val    = units_val.lower()

            # include both “edible (liquid)” and “high cbd edible liquid”
            edible_types = {
                "edible (solid)",
                "edible (liquid)",
                "high cbd edible liquid",
                "tincture",
                "topical",
                "capsule",
            }

            # if it's one of our edible types stored in grams, convert to oz
            if product_type in edible_types and units_val in {"g", "grams"} and weight_val is not None:
                weight_val *= 0.03527396195
                units_val = "oz"

            # now build your display string as before
            if weight_val is not None and units_val:
                weight_str   = f"{weight_val:.2f}".rstrip("0").rstrip(".")
                weight_units = f" -\u00A0{weight_str}{units_val}"
            else:
                weight_units = ""

            # Compute the normal weight string:
            try:
                weight_val = float(row.get("Weight*", ""))
            except:
                weight_val = None
            units_val = row.get("Units", "")
            if weight_val is not None and units_val:
                weight_str = f"{weight_val:.2f}".rstrip("0").rstrip(".")
                normal_weight_units = f" -\u00A0{weight_str}{units_val}"
            else:
                normal_weight_units = ""

            # ── NOW OVERRIDE FOR PARAPHERNALIA ─────────────────────────────
            if product_type == "paraphernalia":
                label_data["WeightUnits"] = ""        # hide it completely
            else:
                label_data["WeightUnits"] = normal_weight_units


            if orientation == "inventory":
                label_data["AcceptedDate"] = str(row.get("Accepted Date", ""))
                label_data["Vendor"] = str(row.get("Vendor", ""))
                label_data["Barcode"] = str(row.get("Barcode*", ""))
                label_data["ProductName"] = str(row.get("Product Name*", ""))
                label_data["ProductType"] = str(row.get("Product Type*", ""))
                # NEW KEY: Wrap the "Quantity*" value with a marker to avoid syntax issues.
                label_data["QuantityReceived"] = str(row.get("Quantity*", ""))
                label_data["WeightUnits"] = weight_units
            
        else:
            label_data = {
                "DOH": "",
                "Price": "",
                "Lineage": "",
                "Ratio_or_THC_CBD": "",
                "ProductBrand": "",
                "Ratio": "",
                "Description": "",
                "ProductStrain": "",
                "WeightUnits": ""
            }
            if orientation == "inventory":
                label_data["AcceptedDate"] = ""
                label_data["Vendor"] = ""
                label_data["Barcode"] = ""
                label_data["ProductName"] = ""
                label_data["ProductType"] = ""
                label_data["QuantityReceived"] = ""
        context[f"Label{i+1}"] = label_data

    tpl.render(context)
    buffer = BytesIO()
    tpl.docx.save(buffer)
    
    doc = Document(BytesIO(buffer.getvalue()))
    
    if orientation != "inventory":
        autosize_fields(doc, font_scheme, orientation, scale_factor=fixed_scale)
        apply_conditional_formatting(doc)
        safe_fix_paragraph_spacing(doc)
        remove_extra_spacing(doc)
        clear_cell_margins(doc)
        clear_table_cell_padding(doc)
        
    if orientation == "mini":
        def fully_clear_cell(cell):
            for child in list(cell._tc):
                cell._tc.remove(child)
            set_cell_background(cell, "FFFFFF")
            cell.text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if not cell.text.strip():
                        fully_clear_cell(cell)
    final_buffer = BytesIO()
    doc.save(final_buffer)
    return final_buffer.getvalue()

# ------------------ Run Full Process Functions ------------------
def filter_column(df, column, var):
    filter_val = normalize(var.get())
    if filter_val != "all" and column in df.columns:
        return df[df[column].astype(str).apply(normalize) == filter_val]
    return df

from io import BytesIO
from docx import Document
from docxcompose.composer import Composer
import datetime
import os
import pandas as pd
from tkinter import messagebox

from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ROW_HEIGHT_RULE

def process_name_chunk(args):
    """
    Given a chunk of records and an in‐memory template buffer, create
    a backside page with the same grid dimensions (3.5"x2.5") filled
    with product names.
    """
    chunk, template_buffer = args
    buf = BytesIO(template_buffer.getvalue())
    doc = Document(buf)

    # Grab the first table
    table = doc.tables[0]

    # Disable autofit so our widths stick
    disable_autofit(table)

    # Set every column to 3.5"
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(3.5)

    # Set every row to 2.5" exact
    for row in table.rows:
        row.height = Inches(2.5)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # Now fill in the names (or leave blank)
    total_cells = len(table.rows) * len(table.columns)
    for idx in range(total_cells):
        r, c = divmod(idx, len(table.columns))
        cell = table.cell(r, c)
        if idx < len(chunk):
            cell.text = chunk[idx]["Product Name*"]
        else:
            cell.text = ""

    out = BytesIO()
    doc.save(out)
    return out.getvalue()

def add_vendor_back_pages(master_doc, records):
    from docx import Document

    # load the same 3×3 template you used for fronts
    back_template = resource_path("templates/horizontal.docx")
    for chunk in chunk_records(records, chunk_size=9):
        back_doc = Document(back_template)
        table   = back_doc.tables[0]

        for idx, row in enumerate(chunk):
            r = idx // 3
            c = idx % 3
            vendor = row.get("Vendor", "").strip()
            brand  = row.get("Product Brand", "").strip()
            text   = f"{vendor}\n{brand}" if vendor or brand else ""
            cell   = table.cell(r, c)
            cell.text = text

        # append this “back” doc right after its front
        composer = Composer(master_doc)
        composer.append(back_doc)


# ─── Main generation function ─────────────────────────────────────
def run_full_process_by_group(template_type, group_by_fields=["Lineage", "Product Strain"]):
    import io, os, datetime
    from docx import Document
    from docxcompose.composer import Composer

    # 1) Get & validate file
    file_path = file_entry.get()
    if not file_path:
        messagebox.showerror("Error", "Please select a data file.")
        return

    # 2) Preprocess & reload
    filters = {
        "product_type": product_type_filter_var.get(),
        "lineage":      lineage_filter_var.get(),
        "brand":        product_brand_filter_var.get(),
        "vendor":       vendor_filter_var.get(),
        "weight":       weight_filter_var.get(),
        "strain":       product_strain_filter_var.get()
    }
    prepped = preprocess_excel(file_path, filters)
    global global_df
    global_df = pd.read_excel(prepped, engine="openpyxl")
    df = global_df.copy()

    # 3) Apply dropdown filters
    df = filter_column(df, "Product Type*", product_type_filter_var)
    df = filter_column(df, "Lineage",         lineage_filter_var)
    df = filter_column(df, "Product Brand",   product_brand_filter_var)
    df = filter_column(df, "Vendor",          vendor_filter_var)
    df = filter_column(df, "CombinedWeight",  weight_filter_var)
    df = filter_column(df, "Product Strain",  product_strain_filter_var)

    # 4) Limit to checked Selected Tags
    sel = [n for n,v in selected_tags_vars.items() if v.get()]
    if not sel:
        messagebox.showerror("Error", "No selected tags are checked!")
        return
    df = df[df["Product Name*"].isin(sel)]

    # 5) Choose template & scheme
    if template_type == "horizontal":
        tpl_path    = resource_path("templates/horizontal.docx")
        orientation = "horizontal"
        scheme      = FONT_SCHEME_HORIZONTAL
    else:
        tpl_path    = resource_path("templates/vertical.docx")
        orientation = "vertical"
        scheme      = FONT_SCHEME_VERTICAL

    # 6) Expand to fixed 3×3
    fixed_buf = expand_template_to_3x3_fixed(tpl_path)

    # 7) Render front & back for each chunk
    bytes_list = []
    records = df.to_dict("records")
    for chunk in chunk_records(records, chunk_size=9):
        # front‐side
        front_bytes = process_chunk((chunk, fixed_buf, scheme, orientation, SCALE_FACTOR))
        bytes_list.append(front_bytes)
        # back‐side vendor sheet?
        if print_vendor_back_var.get():
            back_bytes = process_name_chunk((chunk, fixed_buf))
            bytes_list.append(back_bytes)

    if not bytes_list:
        messagebox.showerror("Error", "No documents generated.")
        return

    # 8) Stitch into master_doc
    master_doc = Document(io.BytesIO(bytes_list[0]))
    composer   = Composer(master_doc)
    for b in bytes_list[1:]:
        composer.append(Document(io.BytesIO(b)))

    # 9) Final spacing
    reapply_table_cell_spacing_only(master_doc, spacing_inches=0.03)

    # 10) Save & open
    today = datetime.datetime.now().strftime("%Y%m%d")
    suffix = "_".join(
        p for p in [filters["product_type"], filters["lineage"],
                    filters["brand"], filters["vendor"],
                    filters["weight"], filters["strain"]]
        if p and p!="All"
    ) or "all"
    out = os.path.join(os.path.expanduser("~"), "Downloads",
                       f"{today}_{orientation}_{suffix}_tags.docx")
    master_doc.save(out)
    open_file(out)
    messagebox.showinfo("Success", f"Saved: {out}")


def run_full_process_mini(bypass_tag_filter: bool = False):
    # ── 0.  Mini template + constants ───────────────────────────────
    base_template       = resource_path("templates/mini.docx")
    orientation         = "mini"
    current_font_scheme = FONT_SCHEME_MINI

    # ── 1.  Pick up the user’s file and filters ─────────────────────
    file_path_val = file_entry.get()
    if not file_path_val:
        messagebox.showerror("Error", "Please select a data file.")
        return

    filters = {
        "product_type": product_type_filter_var.get(),
        "lineage":      lineage_filter_var.get(),
        "brand":        product_brand_filter_var.get(),
        "vendor":       vendor_filter_var.get(),
        "weight":       weight_filter_var.get(),
        "strain":       product_strain_filter_var.get()
    }

    # preprocess once – returns XLSX path + DataFrame cached in RAM
    new_excel_file = preprocess_excel(file_path_val, filters)
    global global_df
    global_df = pd.read_excel(new_excel_file, engine="openpyxl")
    df = global_df.copy()

    # ── 2.  Apply dropdown filters & selected‑tag filter ────────────
    df = filter_column(df, "Product Type*", product_type_filter_var)
    df = filter_column(df, "Lineage",         lineage_filter_var)
    df = filter_column(df, "Product Brand",   product_brand_filter_var)
    df = filter_column(df, "Vendor",          vendor_filter_var)
    df = filter_column(df, "CombinedWeight",  weight_filter_var)
    df = filter_column(df, "Product Strain",  product_strain_filter_var)

    if "Price" in df.columns:
        df["Price"] = df["Price"].apply(lambda x: x.lstrip("'") if isinstance(x, str) else x)

    if not bypass_tag_filter:
        selected_names = [n for n, v in selected_tags_vars.items() if v.get()]
        if not selected_names:
            messagebox.showerror("Error", "No selected tags are checked!")
            return
        df = df[df["Product Name*"].isin(selected_names)]

    if df.empty:
        messagebox.showerror("Error", "No records found after filtering.")
        return

    # ── 3.  Build work items for the pool ───────────────────────────
    records   = df.to_dict(orient="records")
    base_buf  = expand_template_to_5x6_fixed_scaled(
                    base_template,
                    num_inputs=len(records),
                    scale_factor=1.0
                )

    def chunk_records_mini(rec, size=30):      # bigger chunks = faster
        for i in range(0, len(rec), size):
            yield rec[i:i+size]

    work_items = [
        (chunk, base_buf, current_font_scheme, orientation, SCALE_FACTOR)
        for chunk in chunk_records_mini(records)
    ]

    # ── 4.  Render in parallel ──────────────────────────────────────
    from concurrent.futures import ProcessPoolExecutor
    with ProcessPoolExecutor(max_workers=os.cpu_count()) as exe:
        docs_bytes = list(exe.map(process_chunk, work_items))

    docs = [Document(BytesIO(b)) for b in docs_bytes if b]
    if not docs:
        messagebox.showerror("Error", "No documents were generated.")
        return

    # ── 5.  Stitch docs and save ────────────────────────────────────
    master_doc = docs[0]
    composer   = Composer(master_doc)
    for sub_doc in docs[1:]:
        composer.append(sub_doc)
    reapply_table_cell_spacing_only(master_doc)
        # after you generate the front‐side pages in master_doc…
    #if print_vendor_back_var.get():
        # your existing code that inserts the matching
        # vendor‐name back‐pages in the same grid
        #add_vendor_back_pages(master_doc, records)

    today = datetime.datetime.today().strftime("%Y-%m-%d")
    safe = lambda v: str(v).replace(" ", "").replace("/", "").replace("-", "").replace("*", "") if v and v != "All" else None
    suffix_parts = [safe(product_type_filter_var.get()),
                    safe(lineage_filter_var.get()),
                    safe(product_brand_filter_var.get()),
                    safe(vendor_filter_var.get()),
                    safe(weight_filter_var.get()),
                    safe(product_strain_filter_var.get())]
    suffix = "_".join(p for p in suffix_parts if p) or "all"

    doc_path = os.path.join(os.path.expanduser("~"), "Downloads", f"{today}_mini_{suffix}_tags.docx")
    master_doc.save(doc_path)
    open_file(doc_path)
    messagebox.showinfo("Success", f"Word file saved as:\n{doc_path}")



def chunk_records_inv(records, chunk_size=4):
    """Yield chunks of records where each chunk is sized for inventory slip (4 records per slip)."""
    for i in range(0, len(records), chunk_size):
        yield records[i:i+chunk_size]

def run_full_process_inventory_slips(bypass_tag_filter: bool = False):
    """
    Generate 2×2 inventory slips.

    Parameters
    ----------
    bypass_tag_filter : bool, default False
        If True the Selected‑Tag list is ignored (used by the JSON helper).
    """
    # 1) grab user file + filters  ──────────────────────────────────
    file_path_val = file_entry.get()
    if not file_path_val:
        messagebox.showerror("Error", "Please select a data file.")
        return
    filters = {
        "product_type": product_type_filter_var.get(),
        "lineage":      lineage_filter_var.get(),
        "brand":        product_brand_filter_var.get(),
        "vendor":       vendor_filter_var.get(),
        "weight":       weight_filter_var.get(),
        "strain":       product_strain_filter_var.get()
    }

    new_excel_file = preprocess_excel(file_path_val, filters)
    global global_df
    global_df = pd.read_excel(new_excel_file, engine="openpyxl")
    df = global_df.copy()

    # 2) dropdown filters  ─────────────────────────────────────────
    df = filter_column(df, "Product Type*",  product_type_filter_var)
    df = filter_column(df, "Lineage",        lineage_filter_var)
    df = filter_column(df, "Product Brand",  product_brand_filter_var)
    df = filter_column(df, "Vendor",         vendor_filter_var)
    df = filter_column(df, "CombinedWeight", weight_filter_var)
    df = filter_column(df, "Product Strain", product_strain_filter_var)

    # 3) selected‑tag filter unless bypassed  ──────────────────────
    if not bypass_tag_filter:
        selected = [n for n, v in selected_tags_vars.items() if v.get()]
        if not selected:
            messagebox.showerror("Error", "No selected tags are checked!")
            return
        df = df[df["Product Name*"].isin(selected)]

    if df.empty:
        messagebox.showerror("Error", "No records found after filtering.")
        return

    # 4) build template once, then pool‑render  ────────────────────
    base_template   = resource_path("templates/inventorySlips.docx")
    current_scheme  = FONT_SCHEME_VERTICAL
    template_buffer = expand_template_to_2x2_inventory_slips(base_template)
    orientation     = "inventory"

    records   = df.to_dict(orient="records")
    work_items = [
        (chunk, template_buffer, current_scheme, orientation, SCALE_FACTOR)
        for chunk in chunk_records_inv(records)
    ]

    from concurrent.futures import ProcessPoolExecutor
    with ProcessPoolExecutor(max_workers=os.cpu_count()) as exe:
        docs_bytes = list(exe.map(process_chunk, work_items))

    docs = [Document(BytesIO(b)) for b in docs_bytes if b]
    if not docs:
        messagebox.showerror("Error", "No documents were generated.")
        return

    # 5) stitch, save, open  ───────────────────────────────────────
    master_doc = docs[0]
    composer   = Composer(master_doc)
    for sub_doc in docs[1:]:
        composer.append(sub_doc)
        reapply_table_cell_spacing_only(master_doc)
            # after you generate the front‐side pages in master_doc…
        if print_vendor_back_var.get():
            # your existing code that inserts the matching
            # vendor‐name back‐pages in the same grid
            add_vendor_back_pages(master_doc, records)

    today = datetime.datetime.today().strftime("%Y-%m-%d")
    safe = lambda v: str(v).replace(" ", "").replace("/", "").replace("-", "").replace("*", "") if v and v != "All" else None
    suffix = "_".join(filter(None, map(safe, (
        product_type_filter_var.get(), lineage_filter_var.get(),
        product_brand_filter_var.get(), vendor_filter_var.get(),
        weight_filter_var.get(),      product_strain_filter_var.get()
    )))) or "all"

    out_path = os.path.join(Path.home(), "Downloads",
                            f"{today}_inventory_{suffix}_slips.docx")
    master_doc.save(out_path)
    open_file(out_path)
    messagebox.showinfo("Success", f"Inventory slips saved as:\n{out_path}")


def export_data_only():
    messagebox.showinfo("Export Data", "Exported data successfully.")


# ------------------ Global Mousewheel Handler ------------------
def global_mousewheel_handler(event):
    """This handler, bound with bind_all, scrolls the active canvas."""
    global current_canvas
    if current_canvas is None:
        return
    system = platform.system()
    if system == "Darwin":
        # Increase factor as needed for macOS trackpads (try 10 if necessary)
        factor = 10
        scroll_units = int(event.delta * factor)
    else:
        scroll_units = int(event.delta / 120)
    current_canvas.yview_scroll(-scroll_units, "units")
    return "break"

# Bind the global mousewheel handler at the root level.
def bind_global_mousewheel(root):
    root.bind_all("<MouseWheel>", global_mousewheel_handler)
    # For Linux:
    root.bind_all("<Button-4>", global_mousewheel_handler)
    root.bind_all("<Button-5>", global_mousewheel_handler)


# ------------------ UI Helper Functions ------------------
dropdown_cache = {}

def on_mousewheel(event, canvas):
    system = platform.system()
    if system == "Darwin":
        # On macOS, trackpad delta values are small;
        # multiply them by a factor (adjust factor as needed)
        factor = 5  # Experiment with this value.
        scroll_units = int(event.delta * factor)
    else:
        scroll_units = int(event.delta / 120)
    canvas.yview_scroll(-scroll_units, "units")
    return "break"

def update_available_tags_all_state_available():
    # Loop through each available tag's BooleanVar and set it to the checkbox state.
    for tag, var in available_tags_vars.items():
        var.set(available_tags_all_var.get())

def select_all_available():
    for var in available_tags_vars.values():
        var.set(True)
        
def build_dropdown_cache(df):
    global dropdown_cache
    cols = ["Product Type*", "Lineage", "Product Brand", "Vendor", "CombinedWeight", "Product Strain"]
    for col in cols:
        if col in df.columns:
            unique_vals = sorted(df[col].dropna().unique().tolist())
            dropdown_cache[col] = unique_vals

def update_option_menu(option_widget, var, colname):
    menu = option_widget["menu"]
    menu.delete(0, "end")
    options = dropdown_cache.get(colname, []).copy()
    options.insert(0, "All")
    for val in options:
        menu.add_command(label=val, command=lambda v=val: var.set(v))

def populate_filter_dropdowns():
    global global_df
    if global_df is None:
        return
    build_dropdown_cache(global_df)
    if "Product Type*" in global_df.columns:
        update_option_menu(product_type_option, product_type_filter_var, "Product Type*")
    if "Lineage" in global_df.columns:
        update_option_menu(lineage_option, lineage_filter_var, "Lineage")
    if "Product Brand" in global_df.columns:
        update_option_menu(product_brand_option, product_brand_filter_var, "Product Brand")
    if "Vendor" in global_df.columns:
        update_option_menu(vendor_option, vendor_filter_var, "Vendor")
    if "CombinedWeight" in global_df.columns:
        update_option_menu(weight_option, weight_filter_var, "CombinedWeight")
    if "Product Strain" in global_df.columns:
        update_option_menu(product_strain_option, product_strain_filter_var, "Product Strain")

def update_all_dropdowns():
    global _UPDATING_FILTERS, global_df
    if _UPDATING_FILTERS:          # already busy → do nothing
        return

    _UPDATING_FILTERS = True
    try:

        # 1. Create a filtered_df that reflects any current filter settings for Product Type, etc.
        filtered_df = global_df.copy()
        
        def filter_df(column, var):
            value = normalize(var.get())
            if value and value != "all" and column in filtered_df.columns:
                return filtered_df[filtered_df[column].astype(str).apply(normalize) == value]
            else:
                return filtered_df

        # Apply filters that matter for your logic—maybe ignoring weight for now or not:
        filtered_df = filter_df("Product Type*", product_type_filter_var)
        filtered_df = filter_df("Lineage", lineage_filter_var)
        filtered_df = filter_df("Product Brand", product_brand_filter_var)
        filtered_df = filter_df("Vendor", vendor_filter_var)
        filtered_df = filter_df("Product Strain", product_strain_filter_var)

        # 2. Update non-weight dropdowns (type, lineage, brand, vendor, strain) from full cache:
        #    (This always shows the complete set of possible values from the entire dataset.)

        # Product Type
        _update_option_menu(product_type_option, product_type_filter_var, "Product Type*",
                            dropdown_cache["Product Type*"])

        # Lineage
        _update_option_menu(lineage_option, lineage_filter_var, "Lineage",
                            dropdown_cache["Lineage"])

        # Product Brand
        _update_option_menu(product_brand_option, product_brand_filter_var, "Product Brand",
                            dropdown_cache["Product Brand"])

        # Vendor
        _update_option_menu(vendor_option, vendor_filter_var, "Vendor",
                            dropdown_cache["Vendor"])

        # Product Strain
        _update_option_menu(product_strain_option, product_strain_filter_var, "Product Strain",
                            dropdown_cache["Product Strain"])

        # 3. Update weight from the actual filtered DataFrame:
        #    (This shows only relevant weight entries based on current filter selections.)
        weight_options = sorted(filtered_df["CombinedWeight"].dropna().unique(), 
                                key=lambda x: extract_float(str(x)))
        _update_option_menu(weight_option, weight_filter_var, "CombinedWeight", weight_options)

        # 4. Finally, refresh the product names or re-populate the tag list.
        populate_product_names()

    finally:
        _UPDATING_FILTERS = False


def _update_option_menu(menu_widget, var, colname, value_list):
    """
    Clears and repopulates the option menu with the provided 'value_list'.
    Inserts 'All' at the start, preserving the existing selection if possible.
    """
    menu = menu_widget["menu"]
    menu.delete(0, "end")

    # Insert 'All' at the front:
    all_values = ["All"] + list(value_list)

    current_selection = var.get()
    if current_selection not in all_values:
        # If the current selection isn't in the new list, reset to 'All'.
        current_selection = "All"

    for val in all_values:
        menu.add_command(label=val, command=lambda v=val: var.set(v))

    # Finally, update var to preserve or default to 'All'.
    var.set(current_selection)


from pathlib import Path

def populate_available_tags(names):
    global available_tags_container, available_tags_vars, placeholder_img

    # clear out any old widgets
    for widget in available_tags_container.winfo_children():
        widget.destroy()
    available_tags_vars.clear()

    # if no names, display the placeholder image
    if not names:
            lbl = tkmod.Label(
                selected_tags_container,
                image=placeholder_img,
                bg="white"
            )
            lbl.image = placeholder_img
            lbl.is_placeholder = True     # ← mark it
            lbl.pack(expand=True, pady=20)
            return

    # otherwise, build your checkbuttons as before
    for name in names:
        var = tkmod.BooleanVar(value=True)
        chk = tkmod.Checkbutton(available_tags_container,
                                text=name,
                                variable=var,
                                bg="white",
                                anchor="w")
        chk.tag_name = name
        chk.pack(fill="x", pady=2)
        available_tags_vars[name] = var



def populate_selected_tags(names):
    global selected_tags_container, selected_tags_vars, placeholder_img

    # make the panel background grey
    selected_tags_container.configure(bg="lightgrey")

    # clear out old widgets
    for widget in selected_tags_container.winfo_children():
        widget.destroy()
    selected_tags_vars.clear()

    # if empty, show enlarged placeholder, centered
    if not names:
        # 3× zoom
        placeholder_large = placeholder_img.zoom(1, 1)

        lbl = tkmod.Label(
            selected_tags_container,
            image=placeholder_large,
            bg="white",
            anchor="center"
        )
        lbl.image = placeholder_large  # keep a reference

        # pack with expand + center anchor
        lbl.pack(expand=True, anchor="center")
        return

    # otherwise, create checkbuttons
    for name in names:
        var = tkmod.BooleanVar(value=True)
        chk = tkmod.Checkbutton(
            selected_tags_container,
            text=name,
            variable=var,
            bg="white",
            anchor="w"
        )
        chk.tag_name = name
        chk.pack(fill="x", pady=2)
        selected_tags_vars[name] = var


# --- New Section: Selected/Available Tags with "Select All" in Selected Tags ---
selected_tags_all_var = None  # Initialize later in main()

def update_selected_tags_all_state():
    global selected_tags_vars, selected_tags_all_var
    for tag, var in selected_tags_vars.items():
        var.set(selected_tags_all_var.get())

def create_selected_header():
    global selected_tags_all_var
    header_frame = tkmod.Frame(selected_tags_container, bg="lightgray")
    header_frame.pack(fill="x", padx=2, pady=2)
    select_all_chk = tkmod.Checkbutton(header_frame,
                                         text="Select All (Selected Tags)",
                                         variable=selected_tags_all_var,
                                         bg="lightgray",
                                         font=("Arial", 12),
                                         anchor="w",
                                         command=update_selected_tags_all_state)
    select_all_chk.pack(side="left", padx=5)

def move_to_selected():
    global available_tags_vars, selected_tags_vars, undo_stack

    # 1) Find checked tags in Available
    moved_tags = [tag for tag, var in available_tags_vars.items() if var.get()]

    # 2) Move each to Selected
    for tag in moved_tags:
        var = available_tags_vars.pop(tag)
        selected_tags_vars[tag] = var

    # 3) Re-render both lists (placeholder logic lives in those functions)
    populate_available_tags(list(available_tags_vars.keys()))
    populate_selected_tags(list(selected_tags_vars.keys()))

    # 4) Add a divider with current filter values if none exists
    divider_exists = any(getattr(w, "is_divider", False)
                         for w in selected_tags_container.winfo_children())
    if not divider_exists:
        filter_values = []
        if vendor_filter_var.get() != "All":
            filter_values.append("Vendor: " + vendor_filter_var.get())
        if product_brand_filter_var.get() != "All":
            filter_values.append("Brand: " + product_brand_filter_var.get())
        if product_type_filter_var.get() != "All":
            filter_values.append("Type: " + product_type_filter_var.get())
        if lineage_filter_var.get() != "All":
            filter_values.append("Lineage: " + lineage_filter_var.get())
        if product_strain_filter_var.get() != "All":
            filter_values.append("Ratio: " + product_strain_filter_var.get())
        if weight_filter_var.get() != "All":
            filter_values.append("Weight: " + weight_filter_var.get())

        if not filter_values:
            filter_values.append("All")

        divider_text = "------- Selected Filter Values: " + ", ".join(filter_values) + " -------"
        header_divider = tkmod.Label(
            selected_tags_container,
            text=divider_text,
            font=("Arial", 10, "italic"),
            fg="blue",
            bg="lightgray"
        )
        header_divider.is_divider = True
        header_divider.pack(
            fill="x",
            pady=2,
            before=selected_tags_container.winfo_children()[0]
        )

    # 5) Record the move for undo
    if moved_tags:
        undo_stack.append(moved_tags)


    divider_exists = any(getattr(widget, "is_divider", False) for widget in selected_tags_container.winfo_children())
    if not divider_exists:
        filter_values = []
        if vendor_filter_var.get() != "All":
            filter_values.append("Vendor: " + vendor_filter_var.get())
        if product_brand_filter_var.get() != "All":
            filter_values.append("Brand: " + product_brand_filter_var.get())
        if product_type_filter_var.get() != "All":
            filter_values.append("Type: " + product_type_filter_var.get())
        if lineage_filter_var.get() != "All":
            filter_values.append("Lineage: " + lineage_filter_var.get())
        if product_strain_filter_var.get() != "All":
            filter_values.append("Ratio: " + product_strain_filter_var.get())
        if weight_filter_var.get() != "All":
            filter_values.append("Weight: " + weight_filter_var.get())

        if not filter_values:
            filter_values.append("All")

        divider_text = "------- Selected Filter Values: " + ", ".join(filter_values) + " -------"
        header_divider = tkmod.Label(selected_tags_container, text=divider_text,
                                     font=("Arial", 10, "italic"), fg="blue", bg="lightgray")
        header_divider.is_divider = True
        header_divider.pack(fill="x", pady=2, before=selected_tags_container.winfo_children()[0])

    if moved_tags:
        undo_stack.append(moved_tags)



def undo_last_move():
    global undo_stack, available_tags_vars, selected_tags_vars, available_tags_container, selected_tags_container
    if not undo_stack:
        messagebox.showinfo("Undo", "No moves to undo.")
        return
    last_move = undo_stack.pop()  # Get the last list of moved tags
    for tag in last_move:
        # If the tag is currently in selected tags, remove it from there
        if tag in selected_tags_vars:
            # Remove corresponding widget from selected tags container
            for widget in selected_tags_container.winfo_children():
                if hasattr(widget, "tag_name") and widget.tag_name == tag:
                    widget.destroy()
            var = selected_tags_vars.pop(tag)
            # Set its value to True so that it remains selected when moved back
            var.set(True)
            # Re-add the tag to the available tags container
            new_chk = tkmod.Checkbutton(available_tags_container, text=tag, variable=var, bg="white", anchor="w")
            new_chk.tag_name = tag
            new_chk.pack(fill="x", pady=2)
            available_tags_vars[tag] = var

def clear_selected_list():
    global selected_tags_container, selected_tags_vars, undo_stack
    if selected_tags_container is None:
        logging.warning("Selected tags container is not initialized.")
        return

    # Iterate over a copy of the child widget list
    for widget in list(selected_tags_container.winfo_children()):
        try:
            widget.destroy()
        except Exception as e:
            logging.error("Error destroying widget in clear_selected_list: %s", e)
    # Clear the dictionaries and undo history
    selected_tags_vars.clear()
    undo_stack.clear()

    # Refresh available product names if necessary.
    try:
        update_all_dropdowns()
    except Exception as e:
        logging.error("Error updating dropdowns after clearing selected: %s", e)



def move_to_available():
    global available_tags_vars, selected_tags_vars, available_tags_container, selected_tags_container
    if selected_tags_container is None:
        messagebox.showerror("Error", "Selected tags container is not defined.")
        return

    to_move = [tag for tag, var in selected_tags_vars.items() if var.get()]
    for tag in to_move:
        # Remove the corresponding widget from the selected container.
        for widget in selected_tags_container.winfo_children():
            if getattr(widget, "tag_name", None) == tag:
                widget.destroy()
                break
        # When moving back, create the BooleanVar with True so that the tag remains checked.
        new_var = tkmod.BooleanVar(value=True)
        chk = tkmod.Checkbutton(available_tags_container, text=tag, variable=new_var, bg="white", anchor="w")
        chk.tag_name = tag
        chk.pack(fill="x", pady=2)
        available_tags_vars[tag] = new_var
        del selected_tags_vars[tag]
    if not selected_tags_vars:
        selected_tags_all_var.set(False)


def move_tag_to_selected(tag):
    global available_tags_vars, selected_tags_vars, available_tags_container, selected_tags_container
    # Find and destroy the widget from the available container
    for widget in available_tags_container.winfo_children():
        if getattr(widget, "tag_name", None) == tag:
            widget.destroy()
            break
    # Pop the BooleanVar from available_tags_vars; if none, create a new one.
    var = available_tags_vars.pop(tag, tkmod.BooleanVar(value=True))
    # Create the checkbutton in the selected container.
    new_chk = tkmod.Checkbutton(selected_tags_container, text=tag, variable=var, bg="lightgray", anchor="w")
    new_chk.tag_name = tag
    new_chk.pack(fill="x", pady=2)
    selected_tags_vars[tag] = var

def move_tag_to_available(tag):
    global available_tags_vars, selected_tags_vars, available_tags_container, selected_tags_container
    # Find and destroy the widget from the selected container.
    for widget in selected_tags_container.winfo_children():
        if getattr(widget, "tag_name", None) == tag:
            widget.destroy()
            break
    # Create a new BooleanVar for available.
    new_var = tkmod.BooleanVar(value=False)
    chk = tkmod.Checkbutton(available_tags_container, text=tag, variable=new_var, bg="white", anchor="w")
    chk.tag_name = tag
    chk.pack(fill="x", pady=2)
    available_tags_vars[tag] = new_var
    if tag in selected_tags_vars:
        del selected_tags_vars[tag]

    btn_undo = tkmod.Button(button_container, text="↩️ Undo", font=("Arial", 16), command=undo_last_move)
    btn_plus.grid(row=0, column=0, pady=5)
    btn_minus.grid(row=1, column=0, pady=5)
    clear_selected_btn.grid(row=2, column=0, pady=5)
    btn_undo.grid(row=3, column=0, pady=5)

def edit_template(template_type):
    """
    Opens the specified template file in the system's default application for editing.
    
    Valid template_type values:
       - 'horizontal'
       - 'vertical'
       - 'mini'
       - 'inventory'
    """
    template_type = template_type.lower()
    if template_type == 'horizontal':
        path = resource_path("templates/horizontal.docx")
    elif template_type == 'vertical':
        path = resource_path("templates/vertical.docx")
    elif template_type == 'mini':
        path = resource_path("templates/mini.docx")
    elif template_type == 'inventory':
        path = resource_path("templates/inventorySlips.docx")
    else:
        messagebox.showerror("Error", f"Unknown template type: {template_type}")
        return
    
    open_file(path)


def populate_product_names(sorted_names=None):
    global available_tags_container, selected_tags_container, available_tags_vars, selected_tags_vars, global_df
    # Preserve names already in selected panel.
    current_selected = set(selected_tags_vars.keys())
    df = global_df.copy()
    if product_type_filter_var.get() != "All":
        df = df[df["Product Type*"] == product_type_filter_var.get()]
    if lineage_filter_var.get() != "All":
        df = df[df["Lineage"] == lineage_filter_var.get()]
    if product_brand_filter_var.get() != "All":
        df = df[df["Product Brand"] == product_brand_filter_var.get()]
    if vendor_filter_var.get() != "All":
        df = df[df["Vendor"] == vendor_filter_var.get()]
    if weight_filter_var.get() != "All":
        df = df[df["CombinedWeight"] == weight_filter_var.get()]
    if product_strain_filter_var.get() != "All":
        df = df[df["Product Strain"] == product_strain_filter_var.get()]
    if sorted_names is None:
        names = sorted(df["Product Name*"].dropna().unique())
    else:
        names = sorted(sorted_names)
    # Remove names already selected:
    names = [name for name in names if name not in current_selected]
    for widget in available_tags_container.winfo_children():
        widget.destroy()
    available_tags_vars.clear()
    for name in names:
        var = tkmod.BooleanVar(value=True)
        chk = tkmod.Checkbutton(available_tags_container, text=name, variable=var, bg="white", anchor="w")
        chk.tag_name = name
        chk.pack(fill="x", pady=2)
        available_tags_vars[name] = var

def sort_products_by(column):
    # Your sorting logic here.
    # For example:
    global global_df
    if global_df is None or column not in global_df.columns:
        return
    filtered_df = global_df.copy()
    # (apply additional filters as needed)
    sorted_df = filtered_df.sort_values(by=column, na_position='last')
    # Refresh the product names (which update available/selected tags)
    populate_product_names(sorted_names=sorted_df["Product Name*"].dropna().unique().tolist())

def set_current_canvas(event, canvas):
    global current_canvas
    current_canvas = canvas

def clear_current_canvas(event):
    global current_canvas
    current_canvas = None

def change_lineage():
    import os, datetime, webbrowser
    from concurrent.futures import ThreadPoolExecutor
    from tkinter import ttk
    global global_df, selected_tags_vars, root, file_entry
    executor = ThreadPoolExecutor(max_workers=1)

    if global_df is None:
        messagebox.showerror("Error", "No Excel file is loaded.")
        return

    # 1) Capture old lineages
    old_map = {
        name: str(global_df.loc[
            global_df["Product Name*"] == name, "Lineage"
        ].iloc[0]).upper()
        for name in selected_tags_vars
    }

    # 2) Define colors & options (including paraphernalia in pink)
    lineage_map = {
        "SATIVA":        ("(S)",     "#E74C3C"),
        "INDICA":        ("(I)",     "#8E44AD"),
        "HYBRID":        ("(H)",     "#27AE60"),
        "HYBRID/SATIVA": ("(S/H)",   "#E74C3C"),
        "HYBRID/INDICA": ("(I/H)",   "#8E44AD"),
        "CBD":           ("(CBD)",   "#F1C40F"),
        "MIXED":         ("(M)",     "#2C3E50"),
        "PARAPHERNALIA": ("(P)",     "#FF69B4"),
    }
    OPTIONS = list(lineage_map.keys())
    LOG_PATH = os.path.expanduser("~/Downloads/lineage_changes_log.csv")

    # 3) Build popup window
    popup = tkmod.Toplevel(root)
    popup.title("Change Lineage")
    popup.geometry("900x800")
    popup.configure(bg="white")

    # — Scrollable Canvas + InnerFrame —
    list_frame = tkmod.Frame(popup, bg="white")
    list_frame.pack(fill="both", expand=True, padx=10, pady=10)
    canvas = tkmod.Canvas(list_frame, bg="white", highlightthickness=0)
    sb     = tkmod.Scrollbar(list_frame, orient="vertical", command=canvas.yview)
    canvas.configure(yscrollcommand=sb.set)
    canvas.pack(side="left", fill="both", expand=True)
    sb.pack(side="right", fill="y")

    inner_frame = tkmod.Frame(canvas, bg="white")
    window = canvas.create_window((0,0), window=inner_frame, anchor="nw")
    inner_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
    )
    canvas.bind("<Configure>", lambda e: canvas.itemconfig(window, width=e.width))

    # Enable scroll tracking
    canvas.bind("<Enter>", lambda e: set_current_canvas(e, canvas))
    canvas.bind("<Leave>", lambda e: clear_current_canvas(e, canvas))
    canvas.bind("<MouseWheel>", lambda e: on_mousewheel(e, canvas))
    canvas.bind("<Button-4>",    lambda e: on_mousewheel(e, canvas))
    canvas.bind("<Button-5>",    lambda e: on_mousewheel(e, canvas))
    inner_frame.bind("<MouseWheel>", lambda e: on_mousewheel(e, canvas))
    inner_frame.bind("<Button-4>",    lambda e: on_mousewheel(e, canvas))
    inner_frame.bind("<Button-5>",    lambda e: on_mousewheel(e, canvas))

    # Track user selections here
    popup_vars = {}

    # 4) Populate each row: shaded label + Combobox
    for name in sorted(selected_tags_vars):
        old_lin = old_map[name]
        # special case: paraphernalia always shows PARAPHERNALIA
        prod_type = str(global_df.loc[
            global_df["Product Name*"] == name, "Product Type*"
        ].iloc[0]).strip().lower()
        if prod_type == "paraphernalia":
            old_lin = "PARAPHERNALIA"

        abbr, bg = lineage_map.get(old_lin, ("", "#BDC3C7"))

        row = tkmod.Frame(inner_frame, bg=bg)
        row.pack(fill="x", pady=2)

        # shaded product name + old abbr
        lbl = tkmod.Label(
            row,
            text=f"{name}  {abbr}",
            bg=bg,
            fg="white",
            font=("Arial", 16, "bold"),
            anchor="w",
            padx=6, pady=4
        )
        lbl.pack(side="left", fill="x", expand=True)

        # dropdown for new lineage
        var = tkmod.StringVar(value=old_lin)
        popup_vars[name] = var

        combo = ttk.Combobox(
            row, textvariable=var,
            values=OPTIONS, state="readonly", width=12
        )
        combo.pack(side="right", padx=6, pady=4)

    # 5) Save Changes & Cancel buttons
    btn_frame = tkmod.Frame(popup, bg="white")
    btn_frame.pack(fill="x", pady=10)
    tkmod.Button(
        btn_frame, text="Save Changes", font=("Arial",12,"bold"),
        bg="white", fg="green", padx=10, pady=5,
        command=lambda: _apply()
    ).pack(side="right", padx=10)
    tkmod.Button(
        btn_frame, text="Cancel", font=("Arial",12),
        command=popup.destroy
    ).pack(side="right")

    def _apply():
        # 1) Build updated df in‑memory & log diffs
        df2 = global_df.copy()
        ts  = datetime.datetime.now().isoformat()
        os.makedirs(os.path.dirname(LOG_PATH), exist_ok=True)
        with open(LOG_PATH, "a", encoding="utf-8") as log:
            for name, var in popup_vars.items():
                new_lin = var.get().upper()
                old_lin = old_map[name]
                if new_lin != old_lin:
                    df2.loc[
                        df2["Product Name*"] == name, "Lineage"
                    ] = new_lin
                    if new_lin == "MIXED":
                        df2.loc[
                            df2["Product Name*"] == name, "Product Type*"
                        ] = "Mixed"
                    log.write(f"{ts},{name},{old_lin},{new_lin}\n")

        # 2) Background save & reload
        def save_and_reload():
            nowstr = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            out = os.path.expanduser(f"~/Downloads/{nowstr}_LineageUpdated.xlsx")
            df2.to_excel(out, index=False)
            cleaned = preprocess_excel(out)
            newdf = pd.read_excel(cleaned, engine="openpyxl")
            return out, newdf

        future = executor.submit(save_and_reload)

        def on_done(fut):
            try:
                out_path, newdf = fut.result()
            except Exception as e:
                messagebox.showerror("Error", f"Save/Reload failed: {e}")
                return

            # swap in new DataFrame and refresh UI
            global global_df
            global_df = newdf
            populate_filter_dropdowns()
            populate_product_names()
            file_entry.delete(0, "end")
            file_entry.insert(0, out_path)
            messagebox.showinfo(
                "Done",
                f"Saved to:\n{out_path}\n\n"
                f"Log updated at:\n{LOG_PATH}"
            )
            popup.destroy()

        future.add_done_callback(lambda f: root.after_idle(lambda: on_done(f)))

    # Make the popup modal
    popup.grab_set()
    popup.wait_window()


import tkinter as tkmod
from tkinter import ttk, font, colorchooser

def launch_edit_template():
    top = tkmod.Toplevel(root)
    top.title("Edit Template & Font Settings")
    top.geometry("600x500")
    
    nb = ttk.Notebook(top)
    nb.pack(fill="both", expand=True, padx=10, pady=10)
    
    # We'll store settings here:
    font_settings = {
        tmpl: {
            "family": tkmod.StringVar(value="Arial"),
            "size":      tkmod.IntVar(value=12),
            "bold":      tkmod.BooleanVar(value=True),   # <— auto-checked
            "italic": tkmod.BooleanVar(value=False),
            "underline": tkmod.BooleanVar(value=False),
            "color":  tkmod.StringVar(value="#000000"),
        }
        for tmpl in ("Horizontal","Vertical","Mini","Inventory")
    }
    
    def make_font_tab(name):
        frm = ttk.Frame(nb)
        nb.add(frm, text=name)
        
        setting = font_settings[name]
        
        # Font Family
        ttk.Label(frm, text="Font Family:").grid(row=0, column=0, sticky="w", pady=5)
        fam_combo = ttk.Combobox(frm, textvariable=setting["family"],
                                 values=sorted(font.families()), width=30)
        fam_combo.grid(row=0, column=1, sticky="w", pady=5)
        
        # Font Size
        ttk.Label(frm, text="Base Font Size (pt):").grid(row=1, column=0, sticky="w")
        size_spin = tkmod.Spinbox(frm, from_=6, to=72, textvariable=setting["size"], width=5)
        size_spin.grid(row=1, column=1, sticky="w")
        
        # Bold / Italic / Underline
        b1 = tkmod.Checkbutton(frm, text="Bold",      variable=setting["bold"])
        b2 = tkmod.Checkbutton(frm, text="Italic",    variable=setting["italic"])
        b3 = tkmod.Checkbutton(frm, text="Underline", variable=setting["underline"])
        b1.grid(row=2, column=0, sticky="w", pady=5)
        b2.grid(row=2, column=1, sticky="w", pady=5)
        b3.grid(row=2, column=2, sticky="w", pady=5)
        
        # Color Picker
        def choose_color():
            col = colorchooser.askcolor(setting["color"].get(), parent=frm)[1]
            if col:
                setting["color"].set(col)
                color_btn.config(bg=col)
        ttk.Label(frm, text="Font Color:").grid(row=3, column=0, sticky="w")
        color_btn = tkmod.Button(frm, text="   ", command=choose_color,
                                 bg=setting["color"].get(), width=3)
        color_btn.grid(row=3, column=1, sticky="w")
        
        # layout tweaks
        for c in range(3):
            frm.columnconfigure(c, weight=1)
        return frm
    
    # create one tab per template
    for tpl in ("Horizontal","Vertical","Mini","Inventory"):
        make_font_tab(tpl)
    
    # at bottom: OK / Cancel
    btn_frame = ttk.Frame(top)
    btn_frame.pack(fill="x", pady=10)
    def on_ok():
        # here you have all the font_settings[...] values
        # e.g. font_settings["Horizontal"]["family"].get(), etc.
        # save them or apply to your template rendering logic
        top.destroy()
    ttk.Button(btn_frame, text="OK",    command=on_ok).pack(side="right", padx=5)
    ttk.Button(btn_frame, text="Cancel",command=top.destroy).pack(side="right")



def show_instructions_popup():
    # Create a new popup window
    popup = tkmod.Toplevel(root)
    popup.title("POSaBit Instructions")
    popup.geometry("1000x800")
    popup.transient(root)
    popup.grab_set()  # Make the popup modal

    # Create a container to hold the text and image side by side
    container = tkmod.Frame(popup, bg="white", padx=10, pady=10)
    container.pack(fill="both", expand=True)

    # Left side: Instruction Text
    text_label = tkmod.Label(
        container,
        text=posabit_instructions,
        fg="gray",
        bg="white",
        font=("Arial", 16),
        justify="left",
        wraplength=350  # Adjust as needed
    )
    text_label.grid(row=0, column=0, sticky="nsew", padx=(0, 10))

    # Right side: The image (assets/step1.png)
    try:
        image_path = os.path.join("assets", "step1.png")
        step1_image = tkmod.PhotoImage(file=image_path)
    except Exception as e:
        logging.error(f"Error loading image at {image_path}: {e}")
        step1_image = None

    if step1_image:
        image_label = tkmod.Label(container, image=step1_image, bg="white")
        image_label.image = step1_image  # keep a reference to avoid garbage collection
        image_label.grid(row=0, column=1, sticky="nsew")
    else:
        image_label = tkmod.Label(container, text="Image not found", bg="white")
        image_label.grid(row=0, column=1, sticky="nsew")

    # Configure grid weights so that columns share available space equally
    container.columnconfigure(0, weight=1)
    container.columnconfigure(1, weight=1)
    container.rowconfigure(0, weight=1)

    # Add a Close button
    close_btn = tkmod.Button(popup, text="Close", font=("Arial", 12), command=popup.destroy)
    close_btn.pack(pady=10)

    # Wait until the popup is closed before returning (modal behavior)
    popup.wait_window()


def simulate_default_upload():
    default_file = get_default_file()  # Make sure this helper function is defined
    if default_file:
        # Set the file_entry widget to the default file path.
        file_entry.delete(0, tkmod.END)
        file_entry.insert(0, default_file)
        label_file.config(text=os.path.basename(default_file))
        logging.debug(f"Default file found: {default_file}")
        try:
            # Process the default file as if it were just uploaded
            cleaned_file = preprocess_excel(default_file)
            logging.debug(f"Preprocessed file: {cleaned_file}")
            global global_df
            global_df = pd.read_excel(cleaned_file, engine="openpyxl")
            logging.debug(f"DataFrame loaded. Columns: {global_df.columns.tolist()}")
            logging.debug(global_df.head())
            populate_filter_dropdowns()
            populate_product_names()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to process default file: {e}")

def get_default_file():
    """
    Searches the user's Downloads folder for Excel files that start with "A Greener Today"
    and returns the most recently modified file.
    """
    downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
    files = [f for f in os.listdir(downloads_dir)
             if f.startswith("A Greener Today") and f.lower().endswith(".xlsx")]
    if files:
        files_full_paths = [os.path.join(downloads_dir, f) for f in files]
        # Return the most recent file (by modification time)
        return max(files_full_paths, key=os.path.getmtime)
    return None

def simulate_default_upload():
    """
    Sets the 'file_entry' widget to the default file path (if found) and processes it as a newly uploaded file.
    """
    default_file = get_default_upload_file()
    if default_file:
        file_entry.delete(0, tkmod.END)
        file_entry.insert(0, default_file)
        label_file.config(text=os.path.basename(default_file))
        logging.debug(f"Default file found: {default_file}")
        try:
            cleaned_file = preprocess_excel(default_file)
            logging.debug(f"Preprocessed file: {cleaned_file}")
            global global_df
            global_df = pd.read_excel(cleaned_file, engine="openpyxl")
            logging.debug(f"DataFrame loaded. Columns: {global_df.columns.tolist()}")
            logging.debug(global_df.head())
            populate_filter_dropdowns()
            populate_product_names()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to process default file: {e}")
    else:
        logging.debug("No default file found.")



def open_inventory_popup():
    """
    Popup window for JSON input. After the user enters a JSON URL and selects the products from the JSON data,
    clicking "OK" will:
      1. Filter the JSON–derived global DataFrame to the selected products.
      2. Save the filtered data to a temporary Excel file.
      3. Update the main file–entry widget.
      4. Immediately run the Inventory Slips process with bypass_tag_filter=True.
    """
    popup = tkmod.Toplevel(root)
    popup.title("Upload JSON for Inventory Slips")
    popup.geometry("600x700")
    popup.configure(bg="white")

    # ----- Top Section: URL Input -----
    input_frame = tkmod.Frame(popup, bg="white")
    input_frame.pack(side="top", fill="x", padx=10, pady=10)

    instructions = tkmod.Label(
        input_frame,
        text="Enter JSON URL to generate Inventory Slips:",
        font=("Arial", 12),
        bg="white",
        wraplength=560
    )
    instructions.pack(side="top", pady=5)

    url_entry = tkmod.Entry(input_frame, font=("Arial", 12), width=50)
    url_entry.insert(0, "https://example.com/yourfile.json")
    url_entry.pack(side="top", pady=5)

    btn_generate = tkmod.Button(
        input_frame,
        text="Generate from URL",
        font=("Arial", 12),
        command=lambda: process_url()
    )
    btn_generate.pack(side="top", pady=5)

    # ----- Middle Section: Product Selection -----
    selection_frame = tkmod.Frame(popup, bg="white")
    selection_frame.pack(side="top", fill="both", expand=True, padx=10, pady=10)

    selection_label = tkmod.Label(
        selection_frame,
        text="Select products to include:",
        font=("Arial", 12),
        bg="white"
    )
    selection_label.pack(side="top", pady=5)

    # Create a canvas with scrollbar for the product list.
    canvas = tkmod.Canvas(selection_frame, bg="white")
    canvas.pack(side="left", fill="both", expand=True)
    scrollbar = tkmod.Scrollbar(selection_frame, orient="vertical", command=canvas.yview)
    scrollbar.pack(side="right", fill="y")
    canvas.configure(yscrollcommand=scrollbar.set)

    # Frame inside the canvas to hold product checkbuttons.
    product_frame = tkmod.Frame(canvas, bg="white")
    canvas.create_window((0, 0), window=product_frame, anchor="nw")
    product_frame.bind("<Configure>", lambda event: canvas.configure(scrollregion=canvas.bbox("all")))

    # Dictionary to hold BooleanVars for product checkbuttons.
    json_product_vars = {}

    def populate_products(df):
        # Clear any existing checkbuttons.
        # To hide widgets
        for widget in product_frame.winfo_children():
            widget.pack_forget()
        # …
        for widget in product_frame.winfo_children():
            widget.pack(fill="x", pady=2)

        product_names = sorted(df["Product Name*"].dropna().unique())
        for name in product_names:
            var = tkmod.BooleanVar(value=True)
            chk = tkmod.Checkbutton(product_frame, text=name, variable=var,
                                    anchor="w", bg="white", font=("Arial", 10))
            chk.pack(fill="x", padx=5, pady=2)
            json_product_vars[name] = var
        product_frame.update_idletasks()
        canvas.config(scrollregion=canvas.bbox("all"))

    # ----- Bottom Section: OK Button -----
    def on_ok():
        selected = [name for name, var in json_product_vars.items() if var.get()]
        if not selected:
            messagebox.showerror("Error", "Please select at least one product.")
            return

        # Filter the JSON-derived global DataFrame to the selected products.
        global global_df
        filtered_df = global_df[global_df["Product Name*"].isin(selected)].copy()
        if filtered_df.empty:
            messagebox.showerror("Error", "No records found for the selected products.")
            return

        # Save the filtered DataFrame to a temporary Excel file.
        today = datetime.datetime.today().strftime("%Y-%m-%d")
        temp_excel_path = os.path.join(os.path.expanduser("~"), "Downloads", f"{today}_json_filtered.xlsx")
        try:
            filtered_df.to_excel(temp_excel_path, index=False, engine="openpyxl")
        except Exception as e:
            messagebox.showerror("Error", f"Error saving filtered Excel: {e}")
            return

        # Update the main file-entry widget.
        file_entry.delete(0, tkmod.END)
        file_entry.insert(0, temp_excel_path)
        label_file.config(text=os.path.basename(temp_excel_path))

        # Immediately run the inventory slip process (bypassing tag selection).
        run_full_process_inventory_slips(bypass_tag_filter=True)

        popup.destroy()

    btn_ok = tkmod.Button(popup, text="OK", font=("Arial", 12), command=on_ok)
    btn_ok.pack(side="bottom", pady=10)

    # ----- JSON URL Processing -----
    
    def process_url():
        import urllib.request, json   
        url = url_entry.get().strip()
        if not url:
            messagebox.showerror("Error", "Please enter a valid URL.")
            return
        try:
            with urllib.request.urlopen(url) as response:
                data = json.loads(response.read().decode())
            # Process the JSON data into a DataFrame.
            raw_date = data.get("est_arrival_at", "")
            clean_date = raw_date.split("T")[0] if "T" in raw_date else raw_date
            transfer_metadata = {
                "Vendor": f"{data.get('from_license_number')} - {data.get('from_license_name')}",
                "Accepted Date": clean_date,
            }
            records = []
            for item in data.get("inventory_transfer_items", []):
                row = {
                    "Product Name*": item.get("product_name"),
                    "Strain": item.get("strain_name"),
                    "Quantity*": item.get("qty"),
                    "Barcode*": item.get("inventory_id"),
                    "Weight*": item.get("unit_weight"),
                    "Product Type*": item.get("inventory_type"),
                }
                row.update(transfer_metadata)
                records.append(row)
            df_json = pd.DataFrame(records)
            df_json["Units"] = "g"
            df_json["Lineage"] = "HYBRID"
            df_json["Product Strain"] = "Mixed"
            df_json["CombinedWeight"] = df_json["Weight*"].astype(str) + df_json["Units"]
            # Optionally save the entire JSON DataFrame.
            today = datetime.datetime.today().strftime("%Y-%m-%d")
            json_excel_path = os.path.join(os.path.expanduser("~"), "Downloads", f"{today}_json_inventory.xlsx")
            df_json.to_excel(json_excel_path, index=False, engine="openpyxl")
            global global_df
            global_df = df_json.copy()
            populate_products(df_json)
        except Exception as e:
            logging.error("Error processing JSON URL: " + str(e))
            messagebox.showerror("Error", f"Failed to process JSON: {e}")

    popup.grab_set()  # Make the popup modal.
    popup.wait_window()


def show_splash(root):
    splash = tkmod.Toplevel()
    splash.title("Loading...")
    splash.overrideredirect(True)  # Remove window borders
    splash.configure(bg="white")

    # Load the splash image
    try:
        splash_image_path = resource_path("assets/splash.png")
        splash_image = tkmod.PhotoImage(file=splash_image_path)
        width, height = splash_image.width(), splash_image.height()
    except Exception as e:
        logging.error(f"Error loading splash image: {e}")
        width, height = 400, 200  # Fallback size if loading fails
        splash_image = None

    # Center the splash screen
    screen_width = splash.winfo_screenwidth()
    screen_height = splash.winfo_screenheight()
    x = (screen_width // 2) - (width // 2)
    y = (screen_height // 2) - (height // 2)
    splash.geometry(f"{width}x{height}+{x}+{y}")

    if splash_image:
        label = tkmod.Label(splash, image=splash_image, bg="white")
        label.image = splash_image  # Keep reference to avoid garbage collection
    else:
        label = ttk.Label(splash, text="Loading, please wait...", font=("Arial", 16), background="white")

    label.pack(expand=True)

    splash.lift()
    splash.attributes("-topmost", True)
    splash.update()

    return splash


class InventorySlipApp:
    def __init__(self, root):
        self.root = root
        root.title("JSON → Inventory Slip Generator")
        root.geometry("600x700")
        self._build_ui()

    def _build_ui(self):
        frm = tk.Frame(self.root, padx=10, pady=10)
        frm.pack(fill="both", expand=True)

        # JSON URL entry
        tk.Label(frm, text="Enter JSON URL:", font=("Arial", 12)).pack(anchor="w")
        self.url_entry = tk.Entry(frm, font=("Arial", 12), width=50)
        self.url_entry.insert(0, "https://example.com/inventory.json")
        self.url_entry.pack(fill="x", pady=(0,10))
        ttk.Button(frm, text="Load JSON", command=self.load_json).pack(pady=(0,20))

        # Scrollable checkbutton list
        tk.Label(frm, text="Select products to include:", font=("Arial", 12)).pack(anchor="w")
        list_frame = tk.Frame(frm)
        list_frame.pack(fill="both", expand=True)

        self.canvas   = tk.Canvas(list_frame, borderwidth=0)
        self.scroll   = tk.Scrollbar(list_frame, orient="vertical", command=self.canvas.yview)
        self.innerfrm = tk.Frame(self.canvas)

        self.innerfrm.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.create_window((0,0), window=self.innerfrm, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scroll.set)

        self.canvas.pack(side="left", fill="both", expand=True)
        self.scroll.pack(side="right", fill="y")

        # Track each row's BooleanVar
        self.product_vars = {}

        ttk.Button(frm, text="Generate 2×2 Slips", command=self.on_save).pack(pady=20)

    def load_json(self):
        url = self.url_entry.get().strip()
        if not url:
            messagebox.showerror("Error", "Please enter a JSON URL.")
            return

        try:
            with urllib.request.urlopen(url) as resp:
                payload = json.loads(resp.read().decode())
        except Exception as e:
            messagebox.showerror("Error", f"Failed to fetch JSON:\n{e}")
            return

        items = payload.get("inventory_transfer_items", [])
        vendor_meta = f"{payload.get('from_license_number','')} - {payload.get('from_license_name','')}"

        # Build a DataFrame of exactly what we need
        records = []
        for itm in items:
            raw = itm.get("est_arrival_at", "")
            date = raw.split("T")[0] if "T" in raw else raw
            records.append({
                "Product Name*": itm.get("product_name", ""),
                "Product Type*": itm.get("inventory_type", ""),
                "Quantity*"    : itm.get("qty", ""),
                "Barcode*"     : itm.get("inventory_id", ""),
                "Accepted Date": date,
                "Vendor"       : vendor_meta,
            })
        self.df = pd.DataFrame(records)

        # Clear prior checkboxes
        for w in self.innerfrm.winfo_children():
            w.destroy()
        self.product_vars.clear()

        # Create a Checkbutton per product
        for idx, row in self.df.iterrows():
            var = tk.BooleanVar(value=True)
            cb  = tk.Checkbutton(
                self.innerfrm,
                text=row["Product Name*"],
                variable=var,
                anchor="w",
                font=("Arial", 11),
                wraplength=500,
                justify="left"
            )
            cb.pack(fill="x", pady=2, padx=5)
            self.product_vars[idx] = var

    def on_save(self):
        # Gather selected records
        sel_idxs = [i for i,v in self.product_vars.items() if v.get()]
        if not sel_idxs:
            messagebox.showerror("Error", "No products selected.")
            return

        recs = self.df.loc[sel_idxs].to_dict("records")
        all_bytes = [ expand_to_2x2(chunk, TEMPLATE_PATH)
                      for chunk in chunk_records(recs, 4) ]

        if not all_bytes:
            messagebox.showerror("Error", "No slips generated.")
            return

        # Stitch together
        master = Document(BytesIO(all_bytes[0]))
        comp   = Composer(master)
        for b in all_bytes[1:]:
            comp.append(Document(BytesIO(b)))

        now     = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        outname = f"{now}_inventory_slips.docx"
        outpath = os.path.join(OUTPUT_DIR, outname)
        master.save(outpath)

        messagebox.showinfo("Done", f"Saved slips to:\n{outpath}")
        try:
            if sys.platform == "darwin":
                os.system(f'open "{outpath}"')
            elif sys.platform == "win32":
                os.startfile(outpath)
            else:
                os.system(f'xdg-open "{outpath}"')
        except Exception:
            pass


# If the log file doesn’t exist yet, create it with a header row
if not os.path.exists(LOG_PATH):
    with open(LOG_PATH, "w", encoding="utf-8") as f:
        f.write("Timestamp,Product,OldLineage,NewLineage\n")

# ------------------ MAIN GUI FUNCTION ------------------
def main():
    global root, vendor_filter_var, product_brand_filter_var, product_type_filter_var
    global lineage_filter_var, product_strain_filter_var, weight_filter_var, quantity_filter_var
    global file_entry, label_file
    global selected_tags_all_var, available_tags_all_var, selected_tags_vars
    global current_canvas, available_tags_container, selected_tags_container
    global placeholder_img
    global print_vendor_back_var


    selected_tags_vars = {}

    root = tkmod.Tk()
    try:
        placeholder_img = tkmod.PhotoImage(
            file=resource_path("assets/placeholder.png"),
            master=root
        )
    except Exception:
        # fallback if resource_path failed
        placeholder_img = tkmod.PhotoImage(
            file="assets/placeholder.png",
            master=root
        )
    root.withdraw()  # Hide main GUI initially until loading is done

    splash = show_splash(root)

    def load_default_file():
        global global_df
        from pathlib import Path
        downloads_dir = Path.home() / "Downloads"
        matching_files = sorted(downloads_dir.glob("A Greener Today*.xlsx"),
                                key=lambda f: f.stat().st_mtime,
                                reverse=True)
        if matching_files:
            default_path = str(matching_files[0])
            global_df = pd.read_excel(default_path, engine="openpyxl")
            logging.debug("Default file loaded: " + default_path)
        else:
            global_df = pd.DataFrame()
            logging.debug("No default file found.")

    # Load file asynchronously
    from concurrent.futures import ThreadPoolExecutor
    executor = ThreadPoolExecutor(max_workers=1)
    future = executor.submit(load_default_file)

    def check_load_complete():
        if future.done():
            splash.destroy()
            root.deiconify()
            setup_gui(root)
        else:
            splash.after(100, check_load_complete)

    def setup_gui(root):
        root.title("AGT Price Tag Transformer")

        # DPI‑aware scaling
        dpi_scaling = root.winfo_pixels('1i') / 72
        root.tk.call('tk', 'scaling', dpi_scaling)

        # Center and scale GUI
        sw = root.winfo_screenwidth()
        sh = root.winfo_screenheight()
        width  = int(sw * 0.95)
        height = int(sh * 0.95)
        x = (sw - width) // 2
        y = (sh - height) // 2
        root.geometry(f"{width}x{height}+{x}+{y}")

        bind_global_mousewheel(root)

    # Build the main GUI frames and widgets
    main_frame = tkmod.Frame(root, bg="#228B22")
    main_frame.pack(fill="both", expand=True)

    # ---------------- Left Frame: Upload and Filters ----------------
    left_frame = tkmod.Frame(main_frame, bg="#228B22", width=200)
    left_frame.pack(side="left", fill="y", padx=10, pady=10)
    left_frame.pack_propagate(False)

    def upload_file():
        file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx"), ("CSV Files", "*.csv")])
        if file_path:
            label_file.config(text=os.path.basename(file_path))
            file_entry.delete(0, tkmod.END)
            file_entry.insert(0, file_path)
            global global_df, selected_tags_vars
            try:
                cleaned_file = preprocess_excel(file_path)
                logging.debug(f"Preprocessed file: {cleaned_file}")
                global_df = pd.read_excel(cleaned_file, engine="openpyxl")
                logging.debug(f"DataFrame loaded. Columns: {global_df.columns.tolist()}")
                logging.debug(global_df.head())
                populate_filter_dropdowns()
                populate_product_names()
                if not selected_tags_vars:
                    selected_tags_vars = {}
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load file: {e}")

    btn_upload = tkmod.Button(left_frame, text="Upload Spreadsheet", command=upload_file,
                               bg="#228B22", font=("Arial", 16), height=4)
    btn_upload.pack(pady=20)

    label_file = tkmod.Label(left_frame, text="No file selected", bg="#228B22", fg="white", font=("Arial", 7))
    label_file.pack(pady=5)

    file_entry = tkmod.Entry(left_frame, bd=0, bg="white", fg="#000716", font=("Arial", 8))
    #file_entry.pack(fill="x", padx=5, pady=5)


    # Pre-populate the file_entry if a default file is found.
    default_file = get_default_upload_file()
    if default_file:
        file_entry.insert(0, default_file)
        label_file.config(text=os.path.basename(default_file))

    filter_defs = [
        ("\nVendor:", "vendor_filter_var", "vendor_option"),
        ("\nBrand:", "product_brand_filter_var", "product_brand_option"),
        ("\nProduct Type:", "product_type_filter_var", "product_type_option"),
        ("\nLineage (S/H/I):", "lineage_filter_var", "lineage_option"),
        ("\nCBD Blend:", "product_strain_filter_var", "product_strain_option"),
        ("\nWeight:", "weight_filter_var", "weight_option")
    ]
    for text, var_name, option_name in filter_defs:
        lbl = tkmod.Label(left_frame, text=text, bg="#228B22", font=("Arial", 16), fg="white")
        lbl.pack(pady=5)
        globals()[var_name] = tkmod.StringVar(left_frame, value="All")
        opt = tkmod.OptionMenu(left_frame, globals()[var_name], "All")
        opt.config(bg="white", width=10)
        opt["menu"].config(bg="white")
        opt.pack(pady=5, fill="x")
        globals()[option_name] = opt

    if platform.system() == "Darwin":
        check_font = ("Arial", 10)
        pady_val = 10
    else:
        check_font = ("Segoe UI", 8)
        pady_val = 10

    quantity_filter_var = tkmod.BooleanVar(value=True)
    quantity_chk = tkmod.Checkbutton(left_frame, text="Only show products with Quantity > 0",
                                     variable=quantity_filter_var, bg="#228B22", font=check_font,
                                     fg="white", selectcolor="#228B22", activebackground="#228B22",
                                     activeforeground="white", highlightthickness=0, anchor="w", padx=5)
    quantity_chk.pack(pady=pady_val, fill="x")

    file_entry = tkmod.Entry(left_frame, bd=0, bg="white", fg="#000716")

    def clear_filters():
        vendor_filter_var.set("All")
        product_brand_filter_var.set("All")
        product_type_filter_var.set("All")
        lineage_filter_var.set("All")
        product_strain_filter_var.set("All")
        weight_filter_var.set("All")
        update_all_dropdowns()
    btn_clear = tkmod.Button(left_frame, text="Clear Filter", command=clear_filters,
                              bg="#228B22", font=("Arial", 16), height=4)
    btn_clear.pack(pady=10, fill="x")

    # ---------------- Center Frame: Tag Panels ----------------
    center_frame = tkmod.Frame(main_frame, bg="green", width=420, height=800)
    center_frame.pack(side="left", padx=10, pady=10, fill="x", expand=True)
    center_frame.pack_propagate(False)

    
    # Container for tag panels and move buttons
    tags_frame = tkmod.Frame(center_frame, bg="green")
    tags_frame.pack(fill="both", expand=True)

    # ---- Available Tags Panel (Left) ----
    available_panel = tkmod.Frame(tags_frame, bg="white", width=400, height=800)
    available_panel.pack(side="left", fill="both", expand=False)
    available_panel.pack_propagate(False)
    available_label = tkmod.Label(available_panel, text="Available Tag List:", bg="white", font=("Arial", 14))
    available_label.pack(pady=5)

    sort_buttons_frame = tkmod.Frame(available_panel, bg="#D3D3D3")
    sort_buttons_frame.pack(fill="x", padx=5, pady=5)

    available_header = tkmod.Frame(available_panel, bg="white")
    available_header.pack(fill="x", padx=5, pady=(0,5))
    available_tags_all_var = tkmod.BooleanVar(root, value=True)
    available_select_all_chk = tkmod.Checkbutton(
        available_header,
        text="Select All (Available)",
        variable=available_tags_all_var,
        bg="white",
        font=("Arial", 12),
        anchor="w",
        command=update_available_tags_all_state_available
    )
    available_select_all_chk.pack(side="left", padx=5)

    available_canvas = tkmod.Canvas(available_panel, bg="white", height=425)
    available_canvas.pack(side="left", fill="both", expand=True)
    available_scrollbar = tkmod.Scrollbar(available_panel, orient="vertical", command=available_canvas.yview)
    available_scrollbar.pack(side="right", fill="y")
    available_canvas.configure(yscrollcommand=available_scrollbar.set)
    available_tags_container = tkmod.Frame(available_canvas, bg="white")
    available_tags_container.bind("<Configure>", lambda event: available_canvas.configure(scrollregion=available_canvas.bbox("all")))
    available_canvas.create_window((0, 0), window=available_tags_container, anchor="nw")
    available_canvas.bind("<Enter>", lambda event: set_current_canvas(event, available_canvas))
    available_canvas.bind("<Leave>", lambda event: clear_current_canvas(event))
    available_canvas.bind("<MouseWheel>", lambda event: on_mousewheel(event, available_canvas))
    available_canvas.bind("<Button-4>", lambda event: available_canvas.yview_scroll(-1, "units"))
    available_canvas.bind("<Button-5>", lambda event: available_canvas.yview_scroll(1, "units"))

    # ---- Move Buttons Panel (Middle) ----
    move_btn_frame = tkmod.Frame(tags_frame, bg="green", width=130, height=800)
    move_btn_frame.pack(side="left", fill="both", padx=5)
    move_btn_frame.pack_propagate(False)
    button_container = tkmod.Frame(move_btn_frame, bg="green")
    button_container.place(relx=0.5, rely=0.5, anchor="center")
    btn_plus = tkmod.Button(button_container, text=">", font=("Arial", 16), command=move_to_selected)
    btn_minus = tkmod.Button(button_container, text="<", font=("Arial", 16), command=move_to_available)
    clear_selected_btn = tkmod.Button(button_container, text="Clear Selected", font=("Arial", 12), command=clear_selected_list)
    btn_undo = tkmod.Button(button_container, text="Undo", font=("Arial", 12), command=undo_last_move)

    # Instructions '?' button directly under Undo
    btn_instructions = tkmod.Button(button_container, text="?", font=("Arial", 16, "bold"),
                                    fg="#228B22", bg="white", relief="raised",
                                    command=show_instructions_popup)

    # Grid layout
    btn_plus.grid(row=0, column=0, pady=15)
    btn_minus.grid(row=1, column=0, pady=15)
    clear_selected_btn.grid(row=2, column=0, pady=15)
    btn_undo.grid(row=3, column=0, pady=15)
    btn_instructions.grid(row=4, column=0, pady=10)  # '?' button placed here

    


    # ---- Selected Tags Panel (Right) ----
    selected_panel = tkmod.Frame(tags_frame, bg="white", width=425, height=800)
    selected_panel.pack(side="left", fill="both", expand=False)
    selected_panel.pack_propagate(False)
    selected_label = tkmod.Label(selected_panel, text="Selected Tag List:", bg="white", font=("Arial", 14))
    selected_label.pack(pady=5)
    selected_header_frame = tkmod.Frame(selected_panel, bg="white")
    selected_header_frame.pack(fill="x", padx=5, pady=5)
    selected_tags_all_var = tkmod.BooleanVar(root, value=True)
    select_all_chk = tkmod.Checkbutton(selected_header_frame,
                                       text="Select All (Selected Tags)",
                                       variable=selected_tags_all_var,
                                       bg="white", font=("Arial", 12),
                                       anchor="w",
                                       command=update_selected_tags_all_state)
    select_all_chk.pack(side="left", padx=5)

    selected_canvas = tkmod.Canvas(selected_panel, bg="white", height=400)
    selected_canvas.pack(side="left", fill="both", expand=True)
    selected_scrollbar = tkmod.Scrollbar(selected_panel, orient="vertical", command=selected_canvas.yview)
    selected_scrollbar.pack(side="right", fill="y")
    selected_canvas.configure(yscrollcommand=selected_scrollbar.set)
    selected_tags_container = tkmod.Frame(selected_canvas, bg="white")
    selected_tags_container.bind("<Configure>", lambda event: selected_canvas.configure(scrollregion=selected_canvas.bbox("all")))
    selected_canvas.create_window((0, 0), window=selected_tags_container, anchor="nw")
    selected_canvas.bind("<Enter>", lambda event: set_current_canvas(event, selected_canvas))
    selected_canvas.bind("<Leave>", lambda event: clear_current_canvas(event))
    selected_canvas.bind("<MouseWheel>", lambda event: on_mousewheel(event, selected_canvas))
    selected_canvas.bind("<Button-4>", lambda event: selected_canvas.yview_scroll(-1, "units"))
    selected_canvas.bind("<Button-5>", lambda event: selected_canvas.yview_scroll(1, "units"))

    # ---------------- Right Frame: Action Buttons ----------------
    right_frame = tkmod.Frame(main_frame, bg="#228B22", width=150)
    right_frame.pack(side="left", fill="y", padx=10, pady=10)
    right_frame.pack_propagate(False)
        # ─── New: Print Vendor to Back checkbox ─────────────────────────
    print_vendor_back_var = tkmod.BooleanVar(value=True)
    vendor_back_chk = tkmod.Checkbutton(
        right_frame,
        text="Print Vendor to Back",
        variable=print_vendor_back_var,
        bg="#228B22",
        fg="white",
        selectcolor="#228B22",
        font=("Arial", 12),
        anchor="w"
    )
    vendor_back_chk.pack(pady=10, fill="x")

       # ─── Scale Factor slider ─────────────────────────
    scale_factor_var = tkmod.DoubleVar(value=1.0)  # default = 1×

    def on_scale_change(val):
        new_scale = scale_factor_var.get()
        # store it somewhere global or pass it into process_chunk
        global SCALE_FACTOR
        SCALE_FACTOR = new_scale
        # (re-run any previews if you like)


    tkmod.Label(
        right_frame,
        text="Font Scale Factor",
        bg="#228B22", fg="white",
        font=("Arial", 12)
    ).pack(pady=(10,0))

        # right after you create your Scale…
    scale_factor_var = tkmod.DoubleVar(value=1.0)
    scale_slider = tkmod.Scale(
        right_frame,
        variable=scale_factor_var,
        from_=0.5, to=2.0,
        resolution=0.05,
        orient="horizontal",
        length=200,
        bg="#228B22",
        fg="white",
        troughcolor="#BBBBBB",
        highlightthickness=0,
        command=on_scale_change
    )
    scale_slider.pack(pady=(0,10))

    # add a Reset button immediately below (or beside) the slider:
    reset_btn = tkmod.Button(
        right_frame,
        text="Reset Scale",
        font=("Arial", 10),
        command=lambda: (
            scale_factor_var.set(1.0),
            on_scale_change(1.0)  # if you want to reapply immediately
        )
    )
    reset_btn.pack(pady=(0,20))






    btn_horizontal = tkmod.Button(
                            right_frame,
                            text="▭ Horizontal Tags",
                            command=lambda: run_full_process_by_group("horizontal"),
                            bg="#228B22", font=("Arial", 16), height=3,
    anchor="w", padx=10
    )
    btn_horizontal.pack(pady=20, fill="x")

    btn_vertical = tkmod.Button(right_frame, text="▯ Vertical Tags",
                                 command=lambda: run_full_process_by_group("vertical"),
                                 bg="#228B22", font=("Arial", 16), height=3)
    btn_vertical.pack(pady=20, fill="x")
    btn_mini = tkmod.Button(right_frame, text="⬜ Mini Tags",
                            command=run_full_process_mini,
                            bg="#228B22", font=("Arial", 16), height=3)
    btn_mini.pack(pady=20, fill="x")

    btn_edit_template = tkmod.Button(right_frame, 
                                    text="🖊️ Edit Template", 
                                    command=launch_edit_template,
                                    bg="#228B22", font=("Arial", 16), height=3)
    btn_edit_template.pack(pady=20, fill="x")
    btn_edit_data = tkmod.Button(right_frame, text="🎨 Fix Lineage",
                                  command=change_lineage,
                                  bg="#228B22", font=("Arial", 16), height=3)
    btn_edit_data.pack(pady=20, fill="x")
    btn_inventory_slips = tkmod.Button(
        right_frame,
        text=" Inventory Slips",
        command=open_inventory_popup,
        bg="white",
        font=("Arial", 16),
        height=3
    )
    btn_inventory_slips.pack(pady=20, fill="x")  # <-- Add this line

    def bind_dropdown_traces():
        vendor_filter_var.trace_add("write", lambda *args: update_all_dropdowns())
        product_brand_filter_var.trace_add("write", lambda *args: update_all_dropdowns())
        product_type_filter_var.trace_add("write", lambda *args: update_all_dropdowns())
        lineage_filter_var.trace_add("write", lambda *args: update_all_dropdowns())
        product_strain_filter_var.trace_add("write", lambda *args: update_all_dropdowns())
        weight_filter_var.trace_add("write", lambda *args: update_all_dropdowns())
    bind_dropdown_traces()

    from pathlib import Path
    downloads_dir = Path.home() / "Downloads"
    # Use glob to get matching files (case-insensitive if needed)
    matching_files = sorted(downloads_dir.glob("A Greener Today*.xlsx"),
                            key=lambda f: f.stat().st_mtime,
                            reverse=True)
    if matching_files:
        default_path = str(matching_files[0])
        # Update the file_entry widget with the default file path
        file_entry.delete(0, tkmod.END)
        file_entry.insert(0, default_path)
        try:
            global_df = pd.read_excel(default_path, engine="openpyxl")
            populate_filter_dropdowns()
            populate_product_names()  # This function should repopulate available tags automatically
            logging.debug("Default file loaded: " + default_path)
        except Exception as e:
            logging.error("Error reading default file: " + str(e))
    else:
        logging.debug("No default file matching 'A Greener Today*.xlsx' found in Downloads.")
        
    simulate_default_upload()
    populate_filter_dropdowns()
    if "Product Name*" not in global_df.columns:
        messagebox.showerror("Missing Column", "'Product Name*' column not found in your uploaded file.")
        return
    populate_product_names()
    # if no tags have been moved yet, show placeholders in both panels:
    populate_selected_tags([]) 
    check_load_complete()

    logging.debug("Entering mainloop")
    root.mainloop()
    logging.debug("After mainloop (should not reach here until window is closed)")

if __name__ == '__main__':
    try:
        main()
    except Exception as e:
        with open("error.log", "w") as f:
            f.write(traceback.format_exc())
        messagebox.showerror("Application Error", "An error occurred. Please check the error.log file for details.")
        sys.exit(1)
